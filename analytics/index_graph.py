import io
import tempfile
from PIL import Image
from io import BytesIO
from collections.abc import Iterable
from django.contrib.auth.decorators import login_required
from django.core.cache import cache
from django.db.models import (
    Case,
    CharField,
    Count,
    Max,
    QuerySet,
    Sum,
    Value,
    When,
)
from django.http import FileResponse, HttpResponseRedirect, JsonResponse
from django.shortcuts import render
from django.utils.safestring import mark_safe
from django.views import View

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import plotly.io as py
from plotly.offline import plot
from plotly.graph_objects import Figure, Pie
from reportlab.lib.pagesizes import letter, landscape
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

from .forms import Filters


from analytics.models import SchoolDetails
from authenticate.models import CustomUser
state_abv_ = 'sc'
    
''' This sends filters in to django querys 
    if all is present in input we remove
    that as it is eqvivalent of all in query'''

# def filter_set(dashboard_filters):
#     filters = []
#     print("filter set", dashboard_filters)

#     for key,val in dashboard_filters.items():
#         if isinstance(val, list):  # if the value is a list, use the __in lookup
#              if all('all' not in s for s in val) and key != "survey_taken_year":
#                 if key == "zipcode":
#                     zipcode_val = list(map(int, val)) 
#                     filters.append((key, zipcode_val))  
#                 else:
#                     filters.append((key, val))  
#              if key == "survey_taken_year":
#                  #convert the list to int
#                  if isinstance(val, list):  # if the value is a list
#                     try:
#                         new_val = list(map(int, val))  # convert all elements in the list to integers
#                         filters.append((key, new_val))  

#                     except ValueError:
#                       pass 
#         else:  # if the value is not a list, just compare it to 'all'
#             if val != 'all':
#                 filters.append((key, val))
#     filters = dict(filters)
#     return filters

def filter_set(dashboard_filters):
    filters = []
    for key,val in dashboard_filters.items():
        if val != 'all' and val != 'All':
            filters.append((key,val))
    filters = dict(filters)
    return filters

def core_exp_percentage(state,year):
    national_sum_1 = dict(SchoolDetails.objects.values_list('unified_sports_component').filter(unified_sports_component=1,survey_taken_year=year).annotate(total = Count('unified_sports_component'))).get(True,0)
    national_sum_2 = dict(SchoolDetails.objects.values_list('youth_leadership_component').filter(youth_leadership_component=1,survey_taken_year=year).annotate(total = Count('youth_leadership_component'))).get(True,0)
    national_sum_3 = dict(SchoolDetails.objects.values_list('whole_school_component').filter(whole_school_component=1,survey_taken_year=year).annotate(total = Count('whole_school_component'))).get(True,0)
    state_sum_1=dict(SchoolDetails.objects.values_list('unified_sports_component').filter(state_abv=state,unified_sports_component=1,survey_taken_year=year).annotate(total = Count('unified_sports_component'))).get(True,0)
    state_sum_2=dict(SchoolDetails.objects.values_list('youth_leadership_component').filter(state_abv=state,youth_leadership_component=1,survey_taken_year=year).annotate(total = Count('youth_leadership_component'))).get(True,0)
    state_sum_3=dict(SchoolDetails.objects.values_list('whole_school_component').filter(state_abv=state,whole_school_component=1,survey_taken_year=year).annotate(total = Count('whole_school_component'))).get(True,0)
    national_sum = national_sum_1 + national_sum_2 + national_sum_3
    state_sum = state_sum_1+state_sum_2 +state_sum_3
    return national_sum,state_sum

def core_experience_data(dashboard_filters,key,range,survey_year=None):
    select_names = {'sports':'unified_sports_component',
                    'leadership':'youth_leadership_component',
                    'whole_school':'whole_school_component'}
    
    filters = filter_set(dashboard_filters)
    filters_national=filter_set(dashboard_filters)
    # print("CHECK", filters)
    if survey_year:
        # print("CHECK1", filters_national)

        filters["survey_taken_year"] = survey_year
        national_sum,state_sum = core_exp_percentage(state=filters['state_abv'],year=survey_year)
    if range == 'national':
        # print("CHECK2", filters_national)

        filters_national["survey_taken_year"] = survey_year
        filters_national.pop('state_abv')
        filters_national.pop('county', None)
        filters = filters_national
    # print("Check 3", filters_national)
    filters_national = add_in_forFilters(filters_national)
    filters = add_in_forFilters(filters)

    # print("Check 4", filters_national)
    if key:#key value for selectnames dictionary
        # print('working this one?')
        data = dict(SchoolDetails.objects.values_list(select_names[key]).filter(**filters).annotate(total = Count(select_names[key])))
        data = data.get(True,0) #considering only participated people
        # print(data)
        if survey_year:
            if filters.get('state_abv',False):
                try:
                    #state percent
                    # print('state_sum:',data,select_names[key],survey_year)
                    return round((data/state_sum)*100,2)
                except ZeroDivisionError:
                    return 0
            else:
                try:
                    #national percent            
                    # print('national_sum:',data,select_names[key],survey_year)
                    return round((data/national_sum)*100,2)
                except ZeroDivisionError:
                    return 0
            
        return data
    else:
        return 0
    
def core_experience_components_data(dashboard_filters,keys,survey_year='2020'):
    select_names = {'sports':'unified_sports_component',
                    'leadership':'youth_leadership_component',
                    'whole_school':'whole_school_component'}
    
    filters = filter_set(dashboard_filters)
    # print("CHECK", filters)
    # if survey_year:
        # print("CHECK1", filters_national)

        # filters["survey_taken_year"] = survey_year
    national_sum,state_sum = core_exp_percentage(state=filters['state_abv'],year=survey_year)
    
    filters = add_in_forFilters(filters)

    if keys and all(key in select_names for key in keys):
        selected_fields = [select_names[key] for key in keys]

        queryset = SchoolDetails.objects.values(*selected_fields).filter(**filters).annotate(total=Count('auto_increment_id'))
        # print("Core:", queryset)
        tuple_data = {tuple(item[field] for field in selected_fields): item['total'] for item in queryset}

        data = {key: value for key, value in tuple_data.items() if all(key)}

        print("Core:", data)
        # data = data.get(True,0) #considering only participated people
        if len(data) == 1:
            numeric_value = next(iter(data.values()))
            print(numeric_value, state_sum)
        else:
            numeric_value = 0  # or handle the case differently if needed
        if survey_year:
                try:
                    #state percent
                    # print(numeric_value/state_sum)
                    # return round((numeric_value/state_sum)*100,2)
                    return numeric_value
                except ZeroDivisionError:
                    return 0
        return data
    else:
        return 0

def survey_years():
    survey_years = SchoolDetails.objects.values_list('survey_taken_year',flat=True).distinct().order_by('survey_taken_year')
    return list(survey_years)


def core_experience_year_raw_data(dashboard_filters):
    survey_year= survey_years()
    response = {'sports':[],'leadership':[],'whole_school':[],'survey_year':survey_year,'state_sports':[],'state_leadership':[],'state_whole_school':[]}
    for year in survey_year:
        response['sports'].append(core_experience_year_raw_data_values(dashboard_filters,'sports',survey_year=year,range='national'))
        response['leadership'].append(core_experience_year_raw_data_values(dashboard_filters,'leadership',survey_year=year,range='national'))
        response['whole_school'].append(core_experience_year_raw_data_values(dashboard_filters,'whole_school',survey_year=year,range='national'))
        response['state_sports'].append(core_experience_year_raw_data_values(dashboard_filters,'sports',survey_year=year,range='state'))
        response['state_leadership'].append(core_experience_year_raw_data_values(dashboard_filters,'leadership',survey_year=year,range='state'))
        response['state_whole_school'].append(core_experience_year_raw_data_values(dashboard_filters,'whole_school',survey_year=year,range='state'))            

    return response

def core_experience_year_raw_data_values(dashboard_filters, key, range, survey_year=None):

    select_names = {
        'sports': 'unified_sports_component',
        'leadership': 'youth_leadership_component',
        'whole_school': 'whole_school_component'
    }
    
    filters = filter_set(dashboard_filters)
    filters_national = filter_set(dashboard_filters)
    
    if survey_year:
        filters["survey_taken_year"] = survey_year
        national_sum, state_sum = core_exp_percentage(state=filters['state_abv'], year=survey_year)
    
    if range == 'national':
        filters_national["survey_taken_year"] = survey_year
        filters_national.pop('state_abv')
        filters_national.pop('county', None)
        filters = filters_national
    
    filters_national = add_in_forFilters(filters_national)
    filters = add_in_forFilters(filters)
    
    if key:  # key value for selectnames dictionary
        data = dict(SchoolDetails.objects.values_list(select_names[key]).filter(**filters).annotate(total=Count(select_names[key])))
        data = data.get(True, 0)  # considering only participated people
        
        return data  # Return the raw value directly
    else:
        return 0


def core_experience_year_data(dashboard_filters):
    survey_year= survey_years()
    response = {'sports':[],'leadership':[],'whole_school':[],'survey_year':survey_year,'state_sports':[],'state_leadership':[],'state_whole_school':[]}
    for year in survey_year:
        response['sports'].append(core_experience_data(dashboard_filters,'sports',survey_year=year,range='national'))
        response['leadership'].append(core_experience_data(dashboard_filters,'leadership',survey_year=year,range='national'))
        response['whole_school'].append(core_experience_data(dashboard_filters,'whole_school',survey_year=year,range='national'))
        response['state_sports'].append(core_experience_data(dashboard_filters,'sports',survey_year=year,range='state'))
        response['state_leadership'].append(core_experience_data(dashboard_filters,'leadership',survey_year=year,range='state'))
        response['state_whole_school'].append(core_experience_data(dashboard_filters,'whole_school',survey_year=year,range='state'))            
    return response


def survey_response_year_data(dashboard_filters):
    # survey_year= survey_years()
    # for year in survey_year:
    #     response['national'].append(core_experience_data(dashboard_filters,'sports',survey_year=year,range='national'))
    #     response['state'].append(core_experience_data(dashboard_filters,'sports',survey_year=year,range='state'))
    # return response


    filters = filter_set(dashboard_filters)
    survey_year = survey_years()
    response = {'national_yes':[0]*len(survey_year), 'national_no':[0]*len(survey_year), 'survey_year':[0]*len(survey_year),'state_yes':[0]*len(survey_year), 'state_no':[0]*len(survey_year)}
    index = 0
    filters = filter_set(dashboard_filters)
    filters.pop('state_abv')
    if "county" in filters:
        filters.pop('county')
    
    for year in survey_year:
        filters['survey_taken_year'] = year
        # print(filters)
        national_data = SchoolDetails.objects.values('survey_taken', 'survey_taken_year').filter(**filters).exclude(school_state='-99').annotate(total=Count('survey_taken')).order_by('survey_taken_year')        #  = SchoolDetails.objects.values('implementation_level','survey_taken_year').filter(**filters).annotate(total = Count('implementation_level')).order_by('implementation_level')
        print("NATIONAL DATA", national_data)
        for val in national_data:
            if val['survey_taken']:
                response["national_yes"][index] = val.get('total',0)
                response["survey_year"][index] = val.get('survey_taken_year',0)
            elif val['survey_taken'] == False:
                response["national_no"][index] = val.get('total',0)


        index+=1

    # print(response)

    index = 0
    filters_state = filter_set(dashboard_filters)
    if "county" in filters_state:
        filters_state.pop('county')
    filters_state = add_in_forFilters(filters_state)
    
    for year in survey_year:
        filters_state['survey_taken_year'] = year
        state_data = SchoolDetails.objects.values('survey_taken', 'survey_taken_year').filter(**filters_state).exclude(school_state='-99').annotate(total=Count('survey_taken')).order_by('survey_taken_year')
        # print("=====TEST1=====", state_data,filters_state)
        print("SURVEY YEAR", state_data, filters_state)
        for val in state_data:
            print("VAL", val)
            if val['survey_taken']:
                response["state_yes"][index] = val.get('total',0)
            elif val['survey_taken'] == False:
                response["state_no"][index] = val.get('total',0)
        index+=1
    

    # # print('implementation_level_data',response)
    national_sum=''
    state_sum=''

    return response

def survey_response_year_percentage_data(response):
    # Calculate the percentage for 'national_yes'
    print("THE RESPONSE", response['state_yes'], response['state_no'])
    national_percentages = [
        round((yes / (yes + no)) * 100, 2) if (yes + no) > 0 else 0
        for yes, no in zip(response['national_yes'], response['national_no'])
    ]
    
    # Calculate the percentage for 'state_yes'
    state_percentages = [
        round((yes / (yes + no)) * 100, 2) if (yes + no) > 0 else 0
        for yes, no in zip(response['state_yes'], response['state_no'])
    ]
    
    # Update the response dictionary with the calculated percentages
    response['national_yes'] = national_percentages
    response['state_yes'] = state_percentages
    
    # Return only the 'national_yes' and 'state_yes' percentage values
    return {
        'national': response['national_yes'],
        'state': response['state_yes'],
        'survey_year': response['survey_year']
    }

def component_level_data(dashboard_filters):
    filters = filter_set(dashboard_filters)
    survey_year = survey_years()
    # print(survey_year)
    response = {"1 Component":[0]*len(survey_year),"2 Components":[0]*len(survey_year),"3 Components":[0]*len(survey_year),"survey_year":[0]*len(survey_year),
                "state_1_component":[0]*len(survey_year),"state_2_component":[0]*len(survey_year),"state_3_component":[0]*len(survey_year)}
    index = 0
    filters = filter_set(dashboard_filters)
    filters.pop('state_abv')
    if "county" in filters:
        filters.pop('county')
    
    for year in survey_year:
        filters['survey_taken_year'] = year
        # print(filters)
        national_data = SchoolDetails.objects.values('num_components','survey_taken_year').filter(**filters).annotate(total = Count('num_components')).order_by('num_components')
        # print("=========TEST=======", national_data,filters)
        for val in national_data:
            if val['num_components'] in [1, 1.00]:
                response["1 Component"][index] = val.get('total',0)
            elif val['num_components'] in [2,2.00]:
                response["2 Components"][index] = val.get('total',0)
            elif val['num_components'] in [3,3.00]:
                response["3 Components"][index] = val.get('total',0)
                
            response["survey_year"][index] = val.get('survey_taken_year',0)
        index+=1

    index = 0
    filters_state = filter_set(dashboard_filters)
    filters_state = add_in_forFilters(filters_state)
    for year in survey_year:
        filters_state['survey_taken_year'] = year
        state_data = SchoolDetails.objects.values('num_components','survey_taken_year').filter(**filters_state).annotate(total = Count('num_components')).order_by('num_components')

        for val in state_data:
            if val['num_components'] in [1,1.00]:
                response["state_1_component"][index] = val.get('total',0)
            elif val['num_components'] in [2,2.00]:
                response["state_2_component"][index] = val.get('total',0)
            elif val['num_components'] in [3,3.00]:
                response["state_3_component"][index] = val.get('total',0)
        index+=1

    # print('implementation_level_data',response)
    national_sum=''
    state_sum=''

    return response


def component_level_percentages(response):
    for i in range(0,len(response['survey_year'])):
            total=0
            for val in ['1 Component','2 Components','3 Components']:
                total += response[val][i]
            for val in ['1 Component','2 Components','3 Components']:
                # print(total)
                try:
                    response[val][i]=round((response[val][i]/total)*100,2)
                except ZeroDivisionError:
                    response[val][i]=0

    for i in range(0,len(response['survey_year'])):
            total=0
            for val in ['state_1_component','state_2_component','state_3_component']:
                total += response[val][i]
            for val in ['state_1_component','state_2_component','state_3_component']:
                # print(total)
                try:
                    response[val][i]=round((response[val][i]/total)*100,2)
                except ZeroDivisionError:
                     response[val][i]=0

    return response
import copy

def component_level(dashboard_filters, getImage = False):
    raw_data=component_level_data(dashboard_filters)
    data = component_level_percentages(copy.deepcopy(raw_data))
    # colors = ['rgba(200, 200, 220, 0.8)', 'rgba(168, 184, 134, 0.8)', 'rgba(170, 180, 200, 0.8)', 'rgba(109, 154, 168, 0.8)', 'rgba(120, 130, 160, 0.8)', 'rgba(140, 138, 173, 0.8)']     # Define your colors here with opacity

    # print("IMPLEVEL:",data)
    df = pd.DataFrame(data)


    # fig = px.line(df, x=df['survey_year'], y=df['emerging'], labels={"survey_year":"year","emerging":"Implementation Level"})#color???
    fig = go.Figure()
    # hovertemplate='Year: %{x}<br>Emerging: %{y}%<br>Raw Data: %{customdata}<extra></extra>'
    fig.add_scatter(x=df['survey_year'],y=df['state_1_component'].round(0), customdata=raw_data['state_1_component'],
        hovertemplate=(
        "%{x}, %{y}%, %{customdata} schools"
         ), name="{0} 1 Component".format(dashboard_filters['state_abv']))   
    fig.add_scatter(x=df['survey_year'],y=df['state_2_component'].round(0),customdata=raw_data['state_2_component'],
        hovertemplate=(
        "%{x}, %{y}%, %{customdata} schools"
         ), name="{0} 2 Components".format((dashboard_filters['state_abv'])))   
    fig.add_scatter(x=df['survey_year'],y=df['state_3_component'].round(0), customdata=raw_data['state_3_component'],
        hovertemplate=(
        "%{x}, %{y}%, %{customdata} schools"
         ), name="{0} 3 Components".format(dashboard_filters['state_abv']))

    fig.add_scatter(x=df['survey_year'],y=df['1 Component'].round(0), customdata=raw_data['1 Component'],
        hovertemplate=(
        "%{x}, %{y}%, %{customdata} schools"
         ), name="1 Component (National)", visible='legendonly',  marker=dict(color="#636EFA"), line={'dash': 'dash'})
    fig.add_scatter(x=df['survey_year'],y=df['2 Components'].round(0), customdata=raw_data['2 Components'],
        hovertemplate=(
        "%{x}, %{y}%, %{customdata} schools"
         ), name="2 Components (National)" , visible='legendonly',  marker=dict(color="#EF553B"), line={'dash': 'dash'})#color="developing")
    fig.add_scatter(x=df['survey_year'],y=df['3 Components'].round(0), customdata=raw_data['3 Components'],
        hovertemplate=(
        "%{x}, %{y}%, %{customdata} schools"
         ), name="3 Components (National)".format(dashboard_filters['state_abv']), visible='legendonly',  marker=dict(color="#00CC96"), line={'dash': 'dash'})#,color="full_implemention")
    title_name = 'Components over time'.format(dashboard_filters['state_abv'])


    fig.update_layout( 
    title={
        'text': title_name,

        "x": 0.5,  # Adjust the x position of the title (0 - left, 0.5 - center, 1 - right)
        "y": 0.9,  # Adjust the y position of the title (0 - bottom, 1 - top)
        "yanchor": "top",  # Anchor point of the title (aligned to the top)

    },
    legend=dict(
        orientation="h",
    ),
    xaxis = dict (
        tickmode='linear',
        tick0 = min(df['survey_year']),
        dtick=1
    ),
    yaxis = dict (
        range=[0,100],
        title="Percentage value (%)",  # Add this line
    ),
    margin=dict(
        t=100
    ),
    hovermode = 'x',

    )
    plot_div = plot(fig, output_type='div', include_plotlyjs=False)
    if getImage:
        img_bytes = io.BytesIO()
        py.write_image(fig, img_bytes, format='png')
        img_bytes.seek(0)
        return img_bytes
    else:
        if (df == 0).all().all():
            return None
        return plot_div
    

# def component_level(dashboard_filters, getImage = False):
#     raw_data=component_level_data(dashboard_filters)
#     print(raw_data, "===========RAW DATA==============")
#     data = component_level_percentages(copy.deepcopy(raw_data))
#     # colors = ['rgba(200, 200, 220, 0.8)', 'rgba(168, 184, 134, 0.8)', 'rgba(170, 180, 200, 0.8)', 'rgba(109, 154, 168, 0.8)', 'rgba(120, 130, 160, 0.8)', 'rgba(140, 138, 173, 0.8)']     # Define your colors here with opacity
#     print(data, "===========data==============")

#     # print("IMPLEVEL:",data)
#     df = pd.DataFrame(data)
#     print(df, "===========df==============")

#     print('SURVEY YEAR',(df == 0).all())
#     print("end")
#     print(raw_data, "CHECK RAW DATA")
#     # fig = px.line(df, x=df['survey_year'], y=df['emerging'], labels={"survey_year":"year","emerging":"Implementation Level"})#color???
#     fig = go.Figure()
#     # hovertemplate='Year: %{x}<br>Emerging: %{y}%<br>Raw Data: %{customdata}<extra></extra>'
#     fig.add_scatter(x=df['survey_year'],y=df['state_1_component'].round(0), customdata=raw_data['state_1_component'],
#         hovertemplate=(
#         "%{x}, %{y}%, %{customdata} schools"
#          ), name="{0} 1 Component".format(dashboard_filters['state_abv']))   
#     fig.add_scatter(x=df['survey_year'],y=df['state_2_component'].round(0),customdata=raw_data['state_2_component'],
#         hovertemplate=(
#         "%{x}, %{y}%, %{customdata} schools"
#          ), name="{0} 2 Components".format((dashboard_filters['state_abv'])))   
#     fig.add_scatter(x=df['survey_year'],y=df['state_3_component'].round(0), customdata=raw_data['state_3_component'],
#         hovertemplate=(
#         "%{x}, %{y}%, %{customdata} schools"
#          ), name="{0} 3 Components".format(dashboard_filters['state_abv']))

#     fig.add_scatter(x=df['survey_year'],y=df['1 Component'].round(0), customdata=raw_data['1 Component'],
#         hovertemplate=(
#         "%{x}, %{y}%, %{customdata} schools"
#          ), name="1 Component (National)", visible='legendonly')
#     fig.add_scatter(x=df['survey_year'],y=df['2 Components'].round(0), customdata=raw_data['2 Components'],
#         hovertemplate=(
#         "%{x}, %{y}%, %{customdata} schools"
#          ), name="2 Components (National)" , visible='legendonly')#color="developing")
#     fig.add_scatter(x=df['survey_year'],y=df['3 Components'].round(0), customdata=raw_data['3 Components'],
#         hovertemplate=(
#         "%{x}, %{y}%, %{customdata} schools"
#          ), name="3 Components (National)".format(dashboard_filters['state_abv']), visible='legendonly')#,color="full_implemention")
#     title_name = 'Components over time'.format(dashboard_filters['state_abv'])


#     fig.update_layout( 
#     title={
#         'text': title_name,

#         "x": 0.5,  # Adjust the x position of the title (0 - left, 0.5 - center, 1 - right)
#         "y": 0.9,  # Adjust the y position of the title (0 - bottom, 1 - top)
#         "yanchor": "top",  # Anchor point of the title (aligned to the top)

#     },
#     legend=dict(
#         orientation="h",
#     ),
#     xaxis = dict (
#         tickmode='linear',
#         tick0 = min(df['survey_year']),
#         dtick=1
#     ),
#     yaxis = dict (
#         range=[0,100],
#         title="Percentage value (%)",  # Add this line
#     ),
#     margin=dict(
#         t=100
#     ),

#     )
#     plot_div = plot(fig, output_type='div', include_plotlyjs=False)
#     if getImage:
#         img_bytes = io.BytesIO()
#         py.write_image(fig, img_bytes, format='png')
#         img_bytes.seek(0)
#         return img_bytes
#     else:
#         if (df == 0).all().all():
#             return None
#         return plot_div
    

   

def add_in_forFilters(new_filters):
    filters = []
    for key,val in new_filters.items():
             if isinstance(val, list):  # if the value is a list, use the __in lookup
                if  key != "locale__startswith":
                    filters.append((f"{key}__in", val))
                else:
                    filters.append((key, val))  
             else:
                filters.append((key, val))  
    filters = dict(filters)
    return filters

def school_suvery_data(dashboard_filters, isAdmin):
    new_filters= filter_set(dashboard_filters)
    # print("PRIIINNTTTT", new_filters)
    # state = CustomUser.objects.values('state').filter(username=request.user)[0]
    if isAdmin:
        new_filters.pop('state_abv',None) # as this graph is for all states we remove state filter for this
    new_filters.pop('county',None)
    # print('Filterssssss',new_filters)
    filters = add_in_forFilters(new_filters)
    
    # print(filters)

    return SchoolDetails.objects.values('school_state','survey_taken').filter(**filters).exclude(school_state='-99').annotate(total = Count('survey_taken')).order_by('school_state')

def school_suvery_school_year_data(dashboard_filters, isAdmin):
    new_filters= filter_set(dashboard_filters)
    # print("PRIIINNTTTTI", new_filters)
    if isAdmin:
        new_filters.pop('state_abv',None) # as this graph is for all states we remove state filter for this
    new_filters.pop('county',None)
    new_filters.pop('survey_taken_year',None)

    # print('FilterssssssI',new_filters)
    filters = add_in_forFilters(new_filters)
    
    # print(filters)
    
    return SchoolDetails.objects.values('school_state','survey_taken', 'survey_taken_year').filter(**filters).exclude(school_state='-99').annotate(total = Count('survey_taken')).order_by('school_state')

def school_suvery_components_data(dashboard_filters, isAdmin):
    new_filters= filter_set(dashboard_filters)
    # print("PRIIINNTTTTI", new_filters)
    if isAdmin:
     new_filters.pop('state_abv',None) # as this graph is for all states we remove state filter for this
    new_filters.pop('county',None)
    # print('FilterssssssI',new_filters)
    filters = add_in_forFilters(new_filters)
    
    # print(filters)
    
    return SchoolDetails.objects.values('school_state','survey_taken', 'num_components').filter(**filters).exclude(school_state='-99').annotate(total = Count('survey_taken')).order_by('school_state')

def school_suvery_locale_data(dashboard_filters, isAdmin):
    new_filters= filter_set(dashboard_filters)
    # print("PRIIINNTTTTI", new_filters)
    if isAdmin:
        new_filters.pop('state_abv',None) # as this graph is for all states we remove state filter for this
    new_filters.pop('county',None)
    # print('FilterssssssI',new_filters)
    filters = add_in_forFilters(new_filters)
    
    print(filters)
    
    school_surveys = SchoolDetails.objects.values('school_state','survey_taken').filter(**filters).exclude(school_state='-99', locale='-99')
    school_surveys = school_surveys.annotate(
        grouped_locale=Case(
            When(locale__startswith='City', then=Value('City')),
            When(locale__startswith='Rural', then=Value('Rural')),
            When(locale__startswith='Suburb', then=Value('Suburb')),
            When(locale__startswith='Town', then=Value('Town')),
            default=Value('Other'),
            output_field=CharField(),
        )
    )

    return school_surveys.values('school_state', 'survey_taken', 'grouped_locale').annotate(total = Count('survey_taken')).order_by('school_state')

def school_suvery_school_data(dashboard_filters, isAdmin):
    new_filters= filter_set(dashboard_filters)
    # print("PRIIINNTTTTI", new_filters)
    if isAdmin:
        new_filters.pop('state_abv',None) # as this graph is for all states we remove state filter for this
    new_filters.pop('county',None)
    # print('FilterssssssI',new_filters)
    filters = add_in_forFilters(new_filters)
    
    # print(filters)
    
    return SchoolDetails.objects.values('school_state','survey_taken','gradeLevel_WithPreschool').filter(**filters).exclude(school_state='-99').annotate(total = Count('survey_taken')).order_by('school_state')

def school_survey_year(dashboard_filters, isAdmin = False):
    school_surveys = school_suvery_school_year_data(dashboard_filters, isAdmin)
    data_json={}
    data_json_value={}

    totals = {}
    # print(school_surveys)
    for val in school_surveys:
        if val['school_state'] not in data_json.keys():
            data_json[val['school_state']] = {}
            data_json_value[val['school_state']] = {}
        if str(val['survey_taken']) not in data_json[val['school_state']].keys():
            data_json[val['school_state']][str(val['survey_taken'])] = {}
            data_json_value[val['school_state']][str(val['survey_taken'])] = {}
            totals[(val['school_state'], str(val['survey_taken']))] = 0

        data_json[val['school_state']][str(val['survey_taken'])][val['survey_taken_year']] = val['total']
        data_json_value[val['school_state']][str(val['survey_taken'])][val['survey_taken_year']] = val['total']
        totals[(val['school_state'], str(val['survey_taken']))] += val['total']
    
    # print("========RESULT========")
    # print(data_json)
    # print(totals)
    for state, surveys in data_json.items():
        for survey, grades in surveys.items():
            for grade, total in grades.items():
                if totals[(state, survey)] != 0:
                    data_json[state][survey][grade] = round((total / totals[(state, survey)]) * 100, 0)
                else:
             # Handle the case where totals[(state, survey)] is zero
                    data_json[state][survey][grade] = 0
                data_json_value[state][survey][grade] = total
                # print(data_json)
            
        school_state = list(data_json.keys())
    
    survey_1 = []
    survey_2 = []
    survey_3 = []
    survey_4 = []
    survey_5 = []
    survey_6 = []
    survey_1_value = []
    survey_2_value = []
    survey_3_value = []
    survey_4_value = []
    survey_5_value = []
    survey_6_value = []
    for state in school_state:
        # Get the data for each group, or default to 0 if the group doesn't exist

        survey_1.append(data_json[state].get('True', {}).get(2018, 0))
        survey_2.append(data_json[state].get('True', {}).get(2019, 0))
        survey_3.append(data_json[state].get('True', {}).get(2020, 0))
        survey_4.append(data_json[state].get('True', {}).get(2021, 0))
        survey_5.append(data_json[state].get('True', {}).get(2022, 0))
        survey_6.append(data_json[state].get('True', {}).get(2023, 0))
        survey_1_value.append(data_json_value[state].get('True', {}).get(2018, 0))
        survey_2_value.append(data_json_value[state].get('True', {}).get(2019, 0))
        survey_3_value.append(data_json_value[state].get('True', {}).get(2020, 0))
        survey_4_value.append(data_json_value[state].get('True', {}).get(2021, 0))
        survey_5_value.append(data_json_value[state].get('True', {}).get(2022, 0))
        survey_6_value.append(data_json_value[state].get('True', {}).get(2023, 0))
        # print(survey_5)
      

    traces = [
    {"x": school_state, "y": survey_1, "name": '2018', "trace": None, "custom_data": survey_1_value},
    {"x": school_state, "y": survey_2, "name": '2019', "trace": None, "custom_data": survey_2_value},
    {"x": school_state, "y": survey_3, "name": '2020', "trace": None, "custom_data": survey_3_value},
    {"x": school_state, "y": survey_4, "name": '2021', "trace": None, "custom_data": survey_4_value},
    {"x": school_state, "y": survey_5, "name": '2022', "trace": None, "custom_data": survey_5_value},
    {"x": school_state, "y": survey_6, "name": '2023', "trace": None, "custom_data": survey_6_value}

    ]

    # Sort the traces based on their total values
    colors = ['rgba(200, 200, 220, 0.8)', 'rgba(168, 184, 134, 0.8)', 'rgba(170, 180, 200, 0.8)', 'rgba(109, 154, 168, 0.8)', 'rgba(120, 130, 160, 0.8)', 'rgba(140, 138, 173, 0.8)']     # Define your colors here with opacity

    # Create the actual traces
    for i, trace in enumerate(traces):
        trace["trace"] = go.Bar(x=trace["x"], y=trace["y"], name=trace["name"], customdata=trace["custom_data"],
            hovertemplate='%{y}%, %{customdata} schools, %{x}', offsetgroup=1,     marker=dict(color=colors[i])  # Use the colors list here
)

    traces.sort(key=lambda x: sum(x["y"]), reverse=True)

    # Calculate the total for each state
    survey_total = [sum(x) for x in zip(survey_1, survey_2, survey_3, survey_4, survey_5, survey_6)]

    # Create the trace for the total

    # Create the figure and add the traces
    # Note: The total trace should be added first so that it's behind the other bars
    year = dashboard_filters['survey_taken_year']
    # print('YEAR',year)
    # years_str = ', '.join(map(str, year))  # convert the array to a string
    # years_minus_2008 = [int(year1) - 2008 for year1 in year]  # subtract 2008 from each year
    # formatted_year_str = ', '.join(map(str, years_minus_2008))
    # print('YEAR1',formatted_year_str)
    # print('YEAR2',years_minus_2008)
    # Create the layout for the plot
    layout = dict(
        title='Year {0} ({1}) State Program response rate compared by Survey Year'.format(int(year)-2008,(str(int(year)-1)+'-'+str(year)[-2:])),
        yaxis=dict(range=[0, 100]),
        barmode='group',
        hovermode= 'x',
    )

    # Create the figure and add the traces
    fig = go.Figure(data=[trace["trace"] for trace in traces], layout=layout)

    # Update the figure layout
    fig.update_layout(width=1400, height=500)

    # Generate the plot
    plot_div = plot(fig, output_type='div', include_plotlyjs=False)

    return plot_div


def school_survey_school_level(dashboard_filters, isAdmin = False):

    school_surveys = school_suvery_school_data(dashboard_filters, isAdmin)
    # print("school survey", school_surveys)
    data_json={}
    # if 'school_state__in' in data_json.keys():
    #     for val in school_surveys:
    #         if val['school_state__in'] not in data_json.keys():
    #             data_json[val['school_state__in']] = {}
    #         #data_json[val['school_state']][val['survey_taken']] = 0
    #         data_json[val['school_state__in']][str(val['survey_taken__in'])]=val['total']
    # else:
    totals = {}
    data_json_value={}

    for val in school_surveys:
        if val['school_state'] not in data_json.keys():
            data_json[val['school_state']] = {}
            data_json_value[val['school_state']] = {}
        if str(val['survey_taken']) not in data_json[val['school_state']].keys():
            data_json[val['school_state']][str(val['survey_taken'])] = {}
            data_json_value[val['school_state']][str(val['survey_taken'])] = {}
            totals[(val['school_state'], str(val['survey_taken']))] = 0

        data_json[val['school_state']][str(val['survey_taken'])][val['gradeLevel_WithPreschool']] = val['total']
        data_json_value[val['school_state']][str(val['survey_taken'])][val['gradeLevel_WithPreschool']] = val['total']
        totals[(val['school_state'], str(val['survey_taken']))] += val['total']

    # Convert to percentages
    for state, surveys in data_json.items():
        for survey, grades in surveys.items():
            for grade, total in grades.items():
                if totals[(state, survey)] != 0:
                    data_json[state][survey][grade] = round((total / totals[(state, survey)]) * 100, 0)
                else:
             # Handle the case where totals[(state, survey)] is zero
                    data_json[state][survey][grade] = 0
                
                # data_json[state][survey][grade] = round((total / totals[(state, survey)]) * 100, 0)
                data_json_value[state][survey][grade] = total 
        school_state = list(data_json.keys())
        # school_state =(list(data_json.keys()))[0]

    # print("school_state", school_state, dashboard_filters, data_json)
    #convert to percentages
    # Initialize lists to store the data for each group
    survey_1 = []
    survey_2 = []
    survey_3 = []
    survey_4 = []
    survey_1_value = []
    survey_2_value = []
    survey_3_value = []
    survey_4_value = []

    # Iterate over each state
    for state in school_state:
        # Get the data for each group, or default to 0 if the group doesn't exist
        survey_1.append(data_json[state].get('True', {}).get('1.00', 0))
        survey_2.append(data_json[state].get('True', {}).get('2.00', 0))
        survey_3.append(data_json[state].get('True', {}).get('3.00', 0))
        survey_4.append(data_json[state].get('True', {}).get('4.00', 0))
        survey_1_value.append(data_json_value[state].get('True', {}).get('1.00', 0))
        survey_2_value.append(data_json_value[state].get('True', {}).get('2.00', 0))
        survey_3_value.append(data_json_value[state].get('True', {}).get('3.00', 0))
        survey_4_value.append(data_json_value[state].get('True', {}).get('4.00', 0))
        # survey_none.append(data_json[state].get('False', {}).get(None, 0))

    # Create the traces for the plot
    # trace_1 = go.Bar(x=school_state, y=survey_1, name='Emerging (Yes)')
    # trace_2 = go.Bar(x=school_state, y=survey_2, name='Developing (Yes)')
    # trace_3 = go.Bar(x=school_state, y=survey_3, name='Full Implementation (Yes)')
    # trace_none = go.Bar(x=school_state, y=survey_none, name='None')

    traces = [
    {"x": school_state, "y": survey_1, "name": 'Elementary', "trace": None, "custom_data": survey_1_value},
    {"x": school_state, "y": survey_2, "name": 'High', "trace": None, "custom_data": survey_2_value},
    {"x": school_state, "y": survey_3, "name": 'Middle', "trace": None, "custom_data": survey_3_value},
    {"x": school_state, "y": survey_4, "name": 'Preschool', "trace": None, "custom_data": survey_4_value}

    ]

    # Sort the traces based on their total values
    colors = ['rgba(200, 200, 220, 0.8)', 'rgba(168, 184, 134, 0.8)', 'rgba(170, 180, 200, 0.8)', 'rgba(109, 154, 168, 0.8)', 'rgba(120, 130, 160, 0.8)', 'rgba(140, 138, 173, 0.8)']     # Define your colors here with opacity

    # Create the actual traces
    for i, trace in enumerate(traces):
        trace["trace"] = go.Bar(x=trace["x"], y=trace["y"], name=trace["name"], customdata=trace["custom_data"],
            hovertemplate='%{y}%, %{customdata} schools, %{x}', offsetgroup=1,     marker=dict(color=colors[i])  # Use the colors list here
)

    traces.sort(key=lambda x: sum(x["y"]), reverse=True)

    # Calculate the total for each state
    survey_total = [sum(x) for x in zip(survey_1, survey_2, survey_3)]

    # Create the trace for the total

    # Create the figure and add the traces
    # Note: The total trace should be added first so that it's behind the other bars
    year = dashboard_filters['survey_taken_year']
    # print('YEAR',year)
    # years_str = ', '.join(map(str, year))  # convert the array to a string
    # years_minus_2008 = [int(year1) - 2008 for year1 in year]  # subtract 2008 from each year
    # formatted_year_str = ', '.join(map(str, years_minus_2008))
    # print('YEAR1',formatted_year_str)
    # print('YEAR2',years_minus_2008)
    # Create the layout for the plot
    layout = dict(
        title='Year {0} ({1}) State Program response rate compared by School Level'.format(int(year)-2008,(str(int(year)-1)+'-'+str(year)[-2:])),
        yaxis=dict(range=[0, 100]),
        barmode='group',
        hovermode= 'x',
    )

    # Create the figure and add the traces
    fig = go.Figure(data=[trace["trace"] for trace in traces], layout=layout)

    # Update the figure layout
    fig.update_layout(width=1400, height=500)

    # Generate the plot
    plot_div = plot(fig, output_type='div', include_plotlyjs=False)

    return plot_div

def school_survey_locale_level(dashboard_filters, isAdmin = False):

    school_surveys = school_suvery_locale_data(dashboard_filters, isAdmin)
    # print("school survey", school_surveys)
    data_json={}
    # if 'school_state__in' in data_json.keys():
    #     for val in school_surveys:
    #         if val['school_state__in'] not in data_json.keys():
    #             data_json[val['school_state__in']] = {}
    #         #data_json[val['school_state']][val['survey_taken']] = 0
    #         data_json[val['school_state__in']][str(val['survey_taken__in'])]=val['total']
    # else:
    totals = {}
    data_json_value = {}

    for val in school_surveys:
        if val['school_state'] not in data_json.keys():
            data_json[val['school_state']] = {}
            data_json_value[val['school_state']] = {}

        if str(val['survey_taken']) not in data_json[val['school_state']].keys():
            data_json[val['school_state']][str(val['survey_taken'])] = {}
            data_json_value[val['school_state']][str(val['survey_taken'])] = {}

            totals[(val['school_state'], str(val['survey_taken']))] = 0

        data_json[val['school_state']][str(val['survey_taken'])][val['grouped_locale']] = val['total']
        data_json_value[val['school_state']][str(val['survey_taken'])][val['grouped_locale']] = val['total']
        totals[(val['school_state'], str(val['survey_taken']))] += val['total']

    # Convert to percentages
    for state, surveys in data_json.items():
        for survey, locales in surveys.items():
            for locale, total in locales.items():
                data_json_value[state][survey][locale] = total 
                if totals.get((state, survey), 0) != 0:
                    data_json[state][survey][locale] = round((total / totals[(state, survey)]) * 100, 0)
                else:
                # Handle the case where totals[(state, survey)] is zero
                    data_json[state][survey][locale] = 0  # or any other appropriate value    
    school_state = list(data_json.keys())
        # school_state =(list(data_json.keys()))[0]

    # print("school_state", school_state, dashboard_filters, data_json)
    #convert to percentages
    # Initialize lists to store the data for each group
    survey_1 = []
    survey_2 = []
    survey_3 = []
    survey_4 = []
    survey_none = []
    survey_1_value = []
    survey_2_value = []
    survey_3_value = [] 
    survey_4_value = [] 

    # Iterate over each state
    for state in school_state:
        # Get the data for each group, or default to 0 if the group doesn't exist
        survey_1.append(data_json[state].get('True', {}).get('City', 0))
        survey_2.append(data_json[state].get('True', {}).get('Rural', 0))
        survey_3.append(data_json[state].get('True', {}).get('Suburb', 0))
        survey_4.append(data_json[state].get('True', {}).get('Town', 0))
        survey_1_value.append(data_json_value[state].get('True', {}).get('City', 0))
        survey_2_value.append(data_json_value[state].get('True', {}).get('Rural', 0))
        survey_3_value.append(data_json_value[state].get('True', {}).get('Suburb', 0))
        survey_4_value.append(data_json_value[state].get('True', {}).get('Town', 0))
        # survey_none.append(data_json[state].get('False', {}).get(None, 0))

    # Create the traces for the plot
    # trace_1 = go.Bar(x=school_state, y=survey_1, name='Emerging (Yes)')
    # trace_2 = go.Bar(x=school_state, y=survey_2, name='Developing (Yes)')
    # trace_3 = go.Bar(x=school_state, y=survey_3, name='Full Implementation (Yes)')
    # trace_none = go.Bar(x=school_state, y=survey_none, name='None')

    traces = [
    {"x": school_state, "y": [round(val) for val in survey_1], "name": 'City', "trace": None, "custom_data": survey_1_value},
    {"x": school_state, "y": [round(val) for val in survey_2], "name": 'Rural', "trace": None, "custom_data": survey_2_value},
    {"x": school_state, "y": [round(val) for val in survey_3], "name": 'Suburb', "trace": None, "custom_data": survey_3_value},
    {"x": school_state, "y": [round(val) for val in survey_4], "name": 'Town', "trace": None, "custom_data": survey_4_value}
]

    # Sort the traces based on their total values
    colors = ['rgba(200, 200, 220, 0.8)', 'rgba(168, 184, 134, 0.8)', 'rgba(170, 180, 200, 0.8)', 'rgba(109, 154, 168, 0.8)', 'rgba(120, 130, 160, 0.8)', 'rgba(140, 138, 173, 0.8)']     # Define your colors here with opacity

    # Create the actual traces
    for i, trace in enumerate(traces):
        trace["trace"] = go.Bar(x=trace["x"], y=trace["y"], name=trace["name"], customdata=trace["custom_data"],
            hovertemplate='%{y}%, %{customdata} schools, %{x}', offsetgroup=1,     marker=dict(color=colors[i])  # Use the colors list here
)


    # Create the trace for the total
    traces.sort(key=lambda x: sum(x["y"]), reverse=True)

    # Create the figure and add the traces
    # Note: The total trace should be added first so that it's behind the other bars
    year = dashboard_filters['survey_taken_year']
    # print('YEAR',year)
    # years_str = ', '.join(map(str, year))  # convert the array to a string
    # years_minus_2008 = [int(year1) - 2008 for year1 in year]  # subtract 2008 from each year
    # formatted_year_str = ', '.join(map(str, years_minus_2008))
    # print('YEAR1',formatted_year_str)
    # print('YEAR2',years_minus_2008)
    # Create the layout for the plot
    layout = dict(
        title='Year {0} ({1}) State Program response rate compared by Locale'.format(int(year)-2008,(str(int(year)-1)+'-'+str(year)[-2:])),
        yaxis=dict(range=[0, 100]),
        barmode='group',
        hovermode= 'x',
    )

    # Create the figure and add the traces
    fig = go.Figure(data=[trace["trace"] for trace in traces], layout=layout)

    # Update the figure layout
    fig.update_layout(width=1400, height=500)

    # Generate the plot
    plot_div = plot(fig, output_type='div', include_plotlyjs=False)

    return plot_div

def school_survey_implementation_level(dashboard_filters, isAdmin = False):

    school_surveys = school_suvery_components_data(dashboard_filters, isAdmin)
    # Assuming data_json is your data
    # for state in states:
    # # Filter data for current state and survey_taken=True
    # state_data = [data for data in survey_data if data['school_state'] == state and data['survey_taken']]

    # # Get total for each implementation level and append to respective list
    # level_1_totals.append(sum(data['total'] for data in state_data if data['implementation_level'] == '1'))
    # level_2_totals.append(sum(data['total'] for data in state_data if data['implementation_level'] == '2'))
    # level_3_totals.append(sum(data['total'] for data in state_data if data['implementation_level'] == '3'))
    data_json={}
    # if 'school_state__in' in data_json.keys():
    #     for val in school_surveys:
    #         if val['school_state__in'] not in data_json.keys():
    #             data_json[val['school_state__in']] = {}
    #         #data_json[val['school_state']][val['survey_taken']] = 0
    #         data_json[val['school_state__in']][str(val['survey_taken__in'])]=val['total']
    # else:
    totals = {}
    data_json_value = {}
    for val in school_surveys:
        if val['school_state'] not in data_json.keys():
            data_json[val['school_state']] = {}
            data_json_value[val['school_state']] = {}
        if str(val['survey_taken']) not in data_json[val['school_state']].keys():
            data_json[val['school_state']][str(val['survey_taken'])] = {}
            data_json_value[val['school_state']][str(val['survey_taken'])] = {}

            totals[(val['school_state'], str(val['survey_taken']))] = 0

        data_json[val['school_state']][str(val['survey_taken'])][val['num_components']] = val['total']
        data_json_value[val['school_state']][str(val['survey_taken'])][val['num_components']] = val['total']
        totals[(val['school_state'], str(val['survey_taken']))] += val['total']

    # Convert to percentages
    for state, surveys in data_json.items():
        for survey, levels in surveys.items():
            for level, total in levels.items():
                data_json_value[state][survey][level] = total
                if totals.get((state, survey), 0) != 0:
                    data_json[state][survey][level] = round((total / totals[(state, survey)]) * 100, 0)
                else:
                # Handle the case where totals[(state, survey)] is zero
                    data_json[state][survey][level] = 0  # or any other appropriate value    

                # data_json[state][survey][level] = round((total / totals[(state, survey)]) * 100, 0)
        # data_json_value = list(data_json_value.keys())
        school_state = list(data_json.keys())
            # school_state =(list(data_json.keys()))[0]

    # print("school_state", school_state, dashboard_filters, data_json)
  
        # survey_true_val = [data_json[i].get('True',0) for i in school_state]

    #convert to percentages
    # Initialize lists to store the data for each group
    survey_1 = []
    survey_2 = []
    survey_3 = []
    survey_1_value = []
    survey_2_value = []
    survey_3_value = []

    survey_none = []

    # Iterate over each state
    for state in school_state:
        # Get the data for each group, or default to 0 if the group doesn't exist
        survey_1.append(data_json[state].get('True', {}).get(1, 0))
        survey_2.append(data_json[state].get('True', {}).get(2, 0))
        survey_3.append(data_json[state].get('True', {}).get(3, 0))
        # survey_none.append(data_json[state].get('False', {}).get(None, 0))
        survey_1_value.append(data_json_value[state].get('True', {}).get(1, 0))
        survey_2_value.append(data_json_value[state].get('True', {}).get(2, 0))
        survey_3_value.append(data_json_value[state].get('True', {}).get(3, 0))


    # Create the traces for the plot
    # trace_1 = go.Bar(x=school_state, y=survey_1, name='Emerging (Yes)')
    # trace_2 = go.Bar(x=school_state, y=survey_2, name='Developing (Yes)')
    # trace_3 = go.Bar(x=school_state, y=survey_3, name='Full Implementation (Yes)')
    # trace_none = go.Bar(x=school_state, y=survey_none, name='None')

    traces = [
    {"x": school_state, "y": survey_1, "name": '1 Component', "trace": None, "custom_data": survey_1_value},
    {"x": school_state, "y": survey_2, "name": '2 Components', "trace": None, "custom_data": survey_2_value},
    {"x": school_state, "y": survey_3, "name": '3 Components', "trace": None, "custom_data": survey_3_value}
    ]

    # Sort the traces based on their total values
    traces.sort(key=lambda x: sum(x["y"]), reverse=True)
    # traces.reverse
    # colors = ['rgba(255, 0, 0, 1)', 'rgba(0, 255, 0, 1)', 'rgba(0, 0, 255, 1)']  # Define a list of colors

    # Create the actual traces
    colors = ['rgba(200, 200, 220, 0.8)', 'rgba(168, 184, 134, 0.8)', 'rgba(170, 180, 200, 0.8)', 'rgba(109, 154, 168, 0.8)', 'rgba(120, 130, 160, 0.8)', 'rgba(140, 138, 173, 0.8)']     # Define your colors here with opacity

    for i, trace in enumerate(traces):
        trace["trace"] = go.Bar(
            x=trace["x"], 
            y=trace["y"], 
            customdata=trace["custom_data"],
            hovertemplate='%{y}%, %{customdata} schools, %{x}',

            name=trace["name"],
            offsetgroup=1, 
            base=i*0.1,
            marker=dict(color=colors[i])  # Use the colors list here

            # width=0.5,
            # marker=dict(
            #     color=colors[i % len(colors)],  # Use the index to select a color
            #     opacity=1
            # )
        )

    # Calculate the total for each state
    survey_total = [sum(x) for x in zip(survey_1, survey_2, survey_3)]

    # Create the trace for the total

    # Create the figure and add the traces
    # Note: The total trace should be added first so that it's behind the other bars
    year = dashboard_filters['survey_taken_year']
    # print('YEAR',year)
    # years_str = ', '.join(map(str, year))  # convert the array to a string
    # years_minus_2008 = [int(year1) - 2008 for year1 in year]  # subtract 2008 from each year
    # formatted_year_str = ', '.join(map(str, years_minus_2008))
    # print('YEAR1',formatted_year_str)
    # print('YEAR2',years_minus_2008)
    # Create the layout for the plot

    layout = dict(
        title='Year {0} ({1}) State Program response rate compared by Implementation Level'.format(int(year)-2008,(str(int(year)-1)+'-'+str(year)[-2:])),
        yaxis=dict(range=[0, 100]),
        barmode="group",
        hovermode= 'x',
        bargap=0.1,  # Add this line
        
    )

    # Create the figure and add the traces
    fig = go.Figure(data=[trace["trace"] for trace in traces], layout=layout)

    # Update the figure layout
    fig.update_layout(width=1400, height=500)

    # Generate the plot
    plot_div = plot(fig, output_type='div', include_plotlyjs=False)

    return plot_div

def school_survey(dashboard_filters, isAdmin, getImage=False):

    school_surveys = school_suvery_data(dashboard_filters, isAdmin)
    # print("school survey", school_surveys)
    data_json={}
    # if 'school_state__in' in data_json.keys():
    #     for val in school_surveys:
    #         if val['school_state__in'] not in data_json.keys():
    #             data_json[val['school_state__in']] = {}
    #         #data_json[val['school_state']][val['survey_taken']] = 0
    #         data_json[val['school_state__in']][str(val['survey_taken__in'])]=val['total']
    # else:
    for val in school_surveys:
        if val['school_state'] not in data_json.keys():
            data_json[val['school_state']] = {}
            #data_json[val['school_state']][val['survey_taken']] = 0
        data_json[val['school_state']][str(val['survey_taken'])]=val['total']
   
    school_state = list(data_json.keys())
        # school_state =(list(data_json.keys()))[0]

    # print("school_state", school_state, dashboard_filters, data_json)

    survey_true_val = [data_json[i].get('True',0) for i in school_state]

    #convert to percentages
    for state in school_state:
        # print(state)
        total = data_json[state].get('True', 0) + data_json[state].get('False', 0)

        if total > 0:
            data_json[state]['True'] = round((data_json[state].get('True', 0) / total) * 100, 2)
            data_json[state]['False'] = round((data_json[state].get('False', 0) / total) * 100, 2)
        else:
        # Handle the case where total is zero
            data_json[state]['True'] = 0
            data_json[state]['False'] = 0
        # total = data_json[state].get('True',0) + data_json[state].get('False',0)
        # # print("HERE IS THE TOTAL", total)
        # data_json[state]['True'] = round((data_json[state].get('True',0)/total)*100,2)
        # data_json[state]['False'] = round((data_json[state].get('False',0)/total)*100,2)

    survey_true = [data_json[i].get('True',0) for i in school_state]
    survey_false =[data_json[i].get('False',0) for i in school_state]
    # print("survey_true", survey_true)
    survey_true = [round(value) for value in survey_true]
    trace = [go.Bar(
        x= school_state,
        y = survey_true,
        customdata=survey_true_val,
        hovertemplate='%{y}%, %{customdata} schools, %{x}',
        name='',
        marker=dict(color='rgba(109, 154, 168, 0.8)'))]
    year = dashboard_filters['survey_taken_year']
    # print('YEAR',year)
    years_str = year  # convert the array to a string
    # years_minus_2008 = [int(year1) - 2008 for year1 in year]  # subtract 2008 from each year
    # formatted_year_str = years_minus_2008
    # print('YEAR1',formatted_year_str)
    # print('YEAR2',years_minus_2008)

    # if isinstance(year, list):
        # year = year[0]    
    layout = dict(
        title='Year {0} ({1}) State Program response rate'.format(int(year)-2008,(str(int(year)-1)+'-'+str(year)[-2:])),
        yaxis = dict(
        title="Percentage Value (%)",
        range=[0, 100]),
        barmode='group',
        hovermode= 'x',
    )

    fig = go.Figure(data=trace, layout=layout)
    fig.update_layout(width=1400, height=500)
    plot_div = plot(fig, output_type='div', include_plotlyjs=False)
    if getImage:
        img_bytes = io.BytesIO()
        py.write_image(fig, img_bytes, format='png')
        img_bytes.seek(0)
        return img_bytes
    else:
        if survey_true:
            return plot_div
        else:
             return None
    
def school_survey_over_time(dashboard_filters, getImage = False):
    # print("core_experience_yearly filters", filters)
    raw_data = survey_response_year_data(dashboard_filters)
    # print(raw_data, "===========RAW DATA==============")

    data = survey_response_year_percentage_data(raw_data.copy())
    # print("CORE_EXP_YEAR:",data)

    # print(data, "===========data==============")

    # print("IMPLEVEL:",data)
    df = pd.DataFrame(data)
    # print(df, "===========df==============")

  
    # print(raw_data, "===========RAW DATA1==============")

    
    # fig = px.line(df, x=df['survey_year'], y=df['emerging'], labels={"survey_year":"year","emerging":"Implementation Level"})#color???
    fig = go.Figure()
    # hovertemplate='Year: %{x}<br>Emerging: %{y}%<br>Raw Data: %{customdata}<extra></extra>'
    fig.add_scatter(
        x=df['survey_year'],
        y=df['state'].round(0),
        name="{0} response".format(dashboard_filters['state_abv']),
        customdata=raw_data['state_yes'],
        hovertemplate=(
        "%{x}, %{y}%, %{customdata} schools"
         ),
    )  
    fig.add_scatter(x=df['survey_year'],y=df['national'], line={'dash': 'dash'},  marker=dict(color="#636EFA"), name="National response", visible='legendonly', customdata=raw_data['national_yes'],
        hovertemplate=(
        "%{x}, %{y}%, %{customdata} schools"
         ) )
    title_name = 'Survey response rate over time'.format(dashboard_filters['state_abv'])

# for i in range(0,len(response['survey_year'])):
#             total=0
#             for val in ['emerging','developing','full_implement']:
#                 total += response[val][i]
#             for val in ['emerging','developing','full_implement']:
#                 # print(total)
#                 try:
#                     response[val][i]=round((response[val][i]/total)*100,2)
#                 except ZeroDivisionError:
#                     response[val][i]=0

#     for i in range(0,len(response['survey_year'])):
#             total=0
#             for val in ['state_emerging','state_developing','state_full_implement']:
#                 total += response[val][i]
#             for val in ['state_emerging','state_developing','state_full_implement']:
    fig.update_layout( 
    title={
        'text': title_name,
        # 'y': 0.95,  # Adjust the y position of the title (0 - bottom, 1 - top)
        # 'x': 0.5  # Adjust the x position of the title (0 - left, 0.5 - center, 1 - right)

        "x": 0.5,  # Adjust the x position of the title (0 - left, 0.5 - center, 1 - right)
        "y": 0.9,  # Adjust the y position of the title (0 - bottom, 1 - top)
        "yanchor": "top",  # Anchor point of the title (aligned to the top)
      

    },
    legend=dict(
        orientation="h",
    ),
    xaxis = dict (
        tickmode='linear',
        tick0 = min(df['survey_year']),
        dtick=1
    ),
     yaxis = dict (
        range=[0,100],
        title="Percentage value (%)",  # Add this line
    ),
    margin=dict(
        t=100
    ),
    hovermode = 'x',

    )

    if getImage:
        img_bytes = io.BytesIO()
        py.write_image(fig, img_bytes, format='png')
        img_bytes.seek(0)
        return img_bytes
    else:
        plot_div = plot(fig, output_type='div', include_plotlyjs=False)
        # print("imp level yoo?")
        if (df == 0).all().all():
            return None
        return plot_div

import numpy as np

def core_experience(dashboard_filters, getImage = False):
    # survey_year=SchoolDetails.objects.aggregate(Max('survey_taken_year'))
    # survey_year = survey_year['survey_taken_year__max']
    # print("core experience")
    filters = filter_set(dashboard_filters)
    # print("after filter", filters)
    sports=core_experience_data(dashboard_filters,'sports',range='state')
    leadership = core_experience_data(dashboard_filters,'leadership',range='state')
    wholeschool = core_experience_data(dashboard_filters,'whole_school',range='state')

    # print(sports,leadership,wholeschool)
    # print("final filters:", filters)
    colors = ['rgba(170, 180, 200, 0.8)', 'rgba(109, 154, 168, 0.8)', 'rgba(120, 130, 160, 0.8)']  # Define your colors here with opacity
    core_exp_df = pd.DataFrame(dict(
        lables = ['Unified Sports','Inclusive Youth Leadership','Whole School Engagement'],
        values = [sports,leadership,wholeschool]
    ))

    filters = add_in_forFilters(filters)
    fig = Figure(data=[Pie(labels=core_exp_df['lables'], 
                        values=core_exp_df['values'], 
                        hovertemplate='%{label}, %{value} schools<extra></extra>'),
                        ])

    # Set the color of the pie chart
    fig.update_traces(marker=dict(colors=colors))

    # Set the title of the plot
    fig.update_layout(legend_itemclick=False, legend_itemdoubleclick=False, title_text='Core Experience implementation in {year} {state_abv}'.format(state_abv=filters['state_abv'],year=('('+str(int(filters['survey_taken_year'])-1)+'-'+str(filters['survey_taken_year'])[-2:]+')')))
    plot_div = plot(fig, output_type='div', include_plotlyjs=False)
    if getImage:
#         fig.update_layout(
#     autosize=True,
#     width=900,  # Adjust as needed
#     height=1024,  # Adjust as needed
# )
        img_bytes = io.BytesIO()
        py.write_image(fig, img_bytes, format='png')
        img_bytes.seek(0)
        return img_bytes

    else:
        if sports or leadership or wholeschool:
            return plot_div
        else:
            return None
        

import matplotlib
matplotlib.use('agg')

import matplotlib.pyplot as plt
# matplotlib.use('agg')
from matplotlib_venn import venn2, venn3

import plotly as py
from plotly.offline import iplot
import plotly.graph_objs as go
import plotly.io as pio
pio.renderers.default = 'iframe'

import scipy

    # L_sets=[set([sports, sportsLeadership, sportsLeaderShipWholeSchool]), set([leadership, leadershipWholeSchool, sportsLeaderShipWholeSchool]), set([wholeschool, sportsWholeschool, sportsLeaderShipWholeSchool])]

# core experience with componenets
def core_experience_components(dashboard_filters, year, getImage = False):
    filters = filter_set(dashboard_filters)
    sports = core_experience_components_data(dashboard_filters, ['sports'], year)
    leadership = core_experience_components_data(dashboard_filters, ['leadership'], year)
    wholeschool = core_experience_components_data(dashboard_filters, ['whole_school'], year)
    sportsLeadership = core_experience_components_data(dashboard_filters, ['sports', 'leadership'], year)
    leadershipWholeSchool = core_experience_components_data(dashboard_filters, ['leadership', 'whole_school'], year)
    sportsWholeschool = core_experience_components_data(dashboard_filters, ['whole_school', 'sports'], year)
    sportsLeaderShipWholeSchool = core_experience_components_data(dashboard_filters, ['leadership', 'whole_school', 'sports'], year)

    total_schools = sports + leadership + wholeschool - sportsLeadership - leadershipWholeSchool - sportsWholeschool + sportsLeaderShipWholeSchool

    # Calculate the percentage values
    sports_percentage = ((sports - sportsLeadership - sportsWholeschool + sportsLeaderShipWholeSchool) / total_schools) * 100
    leadership_percentage = ((leadership - sportsLeadership - leadershipWholeSchool + sportsLeaderShipWholeSchool) / total_schools) * 100
    wholeschool_percentage = ((wholeschool - sportsWholeschool - leadershipWholeSchool + sportsLeaderShipWholeSchool) / total_schools) * 100
    sportsLeadership_percentage = ((sportsLeadership - sportsLeaderShipWholeSchool) / total_schools) * 100
    leadershipWholeSchool_percentage = ((leadershipWholeSchool - sportsLeaderShipWholeSchool) / total_schools) * 100
    sportsWholeschool_percentage = ((sportsWholeschool - sportsLeaderShipWholeSchool)/ total_schools) * 100
    sportsLeaderShipWholeSchool_percentage = (sportsLeaderShipWholeSchool / total_schools) * 100

    L_sets=[([sports, sportsLeadership, sportsLeaderShipWholeSchool]), ([leadership, leadershipWholeSchool, sportsLeaderShipWholeSchool]), ([wholeschool, sportsWholeschool, sportsLeaderShipWholeSchool])]
    L_labels = ['Sports', 'Leadership', 'WholeSchool']
    setsArr = [("Unified Sports", sports), ("Inclusive Youth Leadership", leadership), ("Whole School Engagement", wholeschool)]
    setsOnlyArr = [("Unified Sports Only", sports - sportsLeadership - sportsWholeschool + sportsLeaderShipWholeSchool), ("Inclusive Youth Leadership Only", leadership - sportsLeadership - leadershipWholeSchool + sportsLeaderShipWholeSchool), ("Whole School Engagement Only", wholeschool - sportsWholeschool - leadershipWholeSchool + sportsLeaderShipWholeSchool)]

    subsetsArr = [("Unified Sports", sports - sportsLeadership - sportsWholeschool + sportsLeaderShipWholeSchool),("Inlcusive Youth Leadership", leadership - sportsLeadership - leadershipWholeSchool + sportsLeaderShipWholeSchool) , ("Unified Sports + Inclusive Youth Leadership", sportsLeadership - sportsLeaderShipWholeSchool), ("Whole School Engagement", wholeschool - sportsWholeschool - leadershipWholeSchool + sportsLeaderShipWholeSchool), ("Unified Sports + Whole School Engagement", sportsWholeschool - sportsLeaderShipWholeSchool), ("Inclusive Youth Leadership + Whole School Engagement",leadershipWholeSchool - sportsLeaderShipWholeSchool), ("Unified Sports + Inclusive Youth Leadership + Whole School Engagement", sportsLeaderShipWholeSchool)] 
    subsetsPercentArr = [("Unified Sports", sports_percentage),("Inclusive Youth Leadership", leadership_percentage) , ("Unified Sports + Inclusive Youth Leadership", sportsLeadership_percentage), ("Whole School Engagement", wholeschool_percentage), ("Unified Sports + Inclusive Whole School Engagement", sportsWholeschool_percentage), ("Inclusive Youth Leadership + Whole School Engagement",leadershipWholeSchool_percentage), ("Unified Sports + Inclusive Youth Leadership + Whole School Engagement", sportsLeaderShipWholeSchool_percentage)] 

    n_sets = len(L_sets)
    print("SUBSET ARR", subsetsArr)
    print("L sets", L_sets, "L_LABELS", L_labels )
    
    counts = {
        '100': sports - sportsLeadership - sportsWholeschool + sportsLeaderShipWholeSchool,
        '010': leadership - sportsLeadership - leadershipWholeSchool + sportsLeaderShipWholeSchool,
        '001': wholeschool - sportsWholeschool - leadershipWholeSchool + sportsLeaderShipWholeSchool,
        '110': sportsLeadership - sportsLeaderShipWholeSchool,
        '101': sportsWholeschool - sportsLeaderShipWholeSchool,
        '011': leadershipWholeSchool - sportsLeaderShipWholeSchool,
        '111': sportsLeaderShipWholeSchool
    }
    subsets = (
        counts['100'],
        counts['010'],
        counts['110'],
        counts['001'],
        counts['101'],
        counts['011'],
        counts['111'],
    )
    
    v = venn3(subsets=subsets, set_labels=L_labels)
        # else:
        #     v = venn3(L_sets)
    
    plt.close()
    print("V Centers", v.centers, "V Radii", v.radii)
    
    L_shapes = []
    L_annotation = []
    
    L_color = ['rgba(170, 180, 200, 0.8)', 'rgba(109, 154, 168, 0.8)', 'rgba(168, 184, 134, 0.8)']
        # colors = ['rgba(200, 200, 220, 0.8)', 'rgba(168, 184, 134, 0.8)', 'rgba(170, 180, 200, 0.8)', 'rgba(109, 154, 168, 0.8)', 'rgba(120, 130, 160, 0.8)', 'rgba(140, 138, 173, 0.8)']     # Define your colors here with opacity

    L_x_max = []
    L_y_max = []
    L_x_min = []
    L_y_min = []
    
    centers = [(0, 0.3), (0.6, 0.3), (0.3, -0.17)]
    subCenters = [(0, 0.3), (0.6, 0.3), (0.3, 0.3), (0.3, -0.17), (0.11, 0.04), (0.5, 0.04), (0.3, 0.15)]
    textSize = [18, 18, 15, 18, 15, 15, 15]
    radii = [0.4, 0.4, 0.4]
    
    for i in range(0, n_sets):
        shape = go.layout.Shape(
            type="circle",
            xref="x",
            yref="y",
            # x0=v.centers[i][0] - v.radii[i],
            # y0=v.centers[i][1] - v.radii[i],
            # x1=v.centers[i][0] + v.radii[i],
            # y1=v.centers[i][1] + v.radii[i],
            x0=centers[i][0] - radii[i],
            y0=centers[i][1] - radii[i],
            x1=centers[i][0] + radii[i],
            y1=centers[i][1] + radii[i],
            fillcolor=L_color[i],
            line_color=L_color[i],
            opacity=1,
            # legendgroup=f"{setsArr[i][0]}"

        )
        
        L_shapes.append(shape)
        print("LABELS:", v.set_labels[i].get_text())
        if i == 0:
            x_offset = -0.3
            y_offset = 0.2
        elif i == 1:
            x_offset = 0.4
            y_offset = 0.2
        else:
            x_offset = 0
            y_offset = -0.45
        
        print("V set labels", v.set_labels[i])
        anno_set_label = go.layout.Annotation(
            xref="x",
            yref="y",
            # x=v.set_labels[i].get_position()[0] + x_offset,
            # y=v.set_labels[i].get_position()[1] + y_offset,
            x=centers[i][0] + x_offset,
            y=centers[i][1] + y_offset,
            text=setsArr[i][0],
            showarrow=False,
            hovertext=f"{setsArr[i][0]}: {setsArr[i][1]} schools <br>{setsOnlyArr[i][0]}: {setsOnlyArr[i][1]} schools",
            font=dict(size=16, color='black', family='Arial')

        )
        
        L_annotation.append(anno_set_label)
        
        # L_x_max.append(v.centers[i][0] + v.radii[i])
        # L_x_min.append(v.centers[i][0] - v.radii[i])
        # L_y_max.append(v.centers[i][1] + v.radii[i])
        # L_y_min.append(v.centers[i][1] - v.radii[i])
        L_x_max.append(centers[i][0] + radii[i])
        L_x_min.append(centers[i][0] - radii[i])
        L_y_max.append(centers[i][1] + radii[i])
        L_y_min.append(centers[i][1] - radii[i])
    
    n_subsets = int(sum([scipy.special.binom(n_sets, i + 1) for i in range(0, n_sets)]))
    print("SUbsets", n_subsets)
    for i in range(0, n_subsets):
        print("SUbsets label",  v.subset_labels[i])

        # if v.subset_labels[i] != None:
        anno_subset_label = go.layout.Annotation(
            xref="x",
            yref="y",
            # x=v.subset_labels[i].get_position()[0],
            # y=v.subset_labels[i].get_position()[1],
            x=subCenters[i][0],  # Adjusting to match the subset positions
            y=subCenters[i][1],
            text=f"{subsetsPercentArr[i][1]:.0f}%",
            showarrow=False,
            hovertext=f"{subsetsArr[i][0]}: {subsetsArr[i][1]} schools",
            font=dict(
            size=textSize[i],  # Set the desired text size here
            color='black',  # Optional: Set the text color
            family='Arial'  # Optional: Set the font family
        )
        )
        
        L_annotation.append(anno_subset_label)
    
    off_set = 0.2
    
    x_max = max(L_x_max) + off_set
    x_min = min(L_x_min) - off_set
    y_max = max(L_y_max) + off_set
    y_min = min(L_y_min) - off_set
    
    p_fig = go.Figure()
    
    p_fig.update_xaxes(
        range=[x_min, x_max],
        showticklabels=False,
        ticklen=0
    )
    
    p_fig.update_yaxes(
        range=[y_min, y_max],
        scaleanchor="x",
        scaleratio=1,
        showticklabels=False,
        ticklen=0
    )
    
    p_fig.update_layout(
        plot_bgcolor='white',
        margin=dict(b=0, l=10, pad=0, r=10, t=40),
        # width=800,
        # height=400,
        shapes=L_shapes,
        annotations=L_annotation,
        title=dict(text='Components Venn Diagram', x=0.28, xanchor='left'),
        

        # title_text='Core Experience implementation in {year} {state_abv}'
    )
    p_fig.update_layout(title_text='Components Venn Diagram in {year} {state_abv}'.format(state_abv=filters['state_abv'],year=('('+str(int(filters['survey_taken_year'])-1)+'-'+str(filters['survey_taken_year'])[-2:]+')')))

    
    for i in range(len(L_shapes)):
        p_fig.add_trace(go.Scatter(
            x=[None],
            y=[None],
            # text=[f"{L_labels[i]}: {len(L_sets[i])}"],
        mode="markers",
        marker=dict(size=15, color=L_color[i]),
        name=f"{setsArr[i][0]}",
        legendgroup=f"{setsArr[i][0]}",
        showlegend=True,
        hoverinfo="text"
        ))
    
    plot_div = plot(p_fig, output_type='div', include_plotlyjs=False)
    
    return plot_div
   


def core_experience_yearly(dashboard_filters, getImage = False):
    filters = filter_set(dashboard_filters)
    raw_data = core_experience_year_raw_data(dashboard_filters)
    # print("core_experience_yearly filters", filters)
    data = core_experience_year_data(dashboard_filters)

    # print("CORE_EXP_YEAR:",data)

    #{'sports': [6370, 3461], 'leadership': [3914, 2560], 'whole_school': [5132, 3361], 'survey_year': [2021, 2022], 'state_sports': [47, 21], 'state_leadership': [18, 12], 'state_whole_school': [26, 18]}
    # print(raw_data, "CORE EPEXRIENCE RAW DATA")
    df = pd.DataFrame(data)
    # print("CORE EXP SURVEY YEAR",(df))
    fig = go.Figure()
    fig.add_scatter(x=df['survey_year'],y=df['state_leadership'].round(0), customdata=raw_data['state_leadership'],
        hovertemplate=(
        "%{x}, %{y}%, %{customdata} schools"
         ), name="{state} Inclusive Youth Leadership".format(state=filters['state_abv']))
    fig.add_scatter(x=df['survey_year'],y=df['state_sports'].round(0), customdata=raw_data['state_sports'],
        hovertemplate=(
        "%{x}, %{y}%, %{customdata} schools"
         ), name="{state} Unified Sports".format(state=filters['state_abv']))
    fig.add_scatter(x=df['survey_year'],y=df['state_whole_school'].round(0), customdata=raw_data['state_whole_school'],
        hovertemplate=(
        "%{x}, %{y}%, %{customdata} schools"
         ), name="{state} School Engagement".format(state=filters['state_abv']))

    fig.add_scatter(x=df['survey_year'],y=df['sports'].round(0), customdata=raw_data['sports'],
        hovertemplate=(
        "%{x}, %{y}%, %{customdata} schools"
         ), name="Unified Sports (National)",  visible='legendonly',marker=dict(color="#636EFA"), line={'dash': 'dash'} )
    fig.add_scatter(x=df['survey_year'],y=df['whole_school'].round(0),customdata=raw_data['whole_school'],
        hovertemplate=(
        "%{x}, %{y}%, %{customdata} schools"
         ),  name="Whole School Engagement (National)",  visible='legendonly', marker=dict(color="#EF553B"), line={'dash': 'dash'})
    fig.add_scatter(x=df['survey_year'],y=df['leadership'].round(0), customdata=raw_data['leadership'],
        hovertemplate=(
        "%{x}, %{y}%, %{customdata} schools"
         ), name="Inclusive Youth Leadership (National)",  visible='legendonly', marker=dict(color="#00CC96"), line={'dash': 'dash'})
    title_name = "Percentage of Core experience implementation over time,<br> National vs {0} State program <br><br>".format(filters['state_abv'])
    
    fig.update_layout( 
        
    title={
        'text': title_name,
        # 'y': 0.95,  # Adjust the y position of the title (0 - bottom, 1 - top)
        # 'x': 0.5,  # Adjust the x position of the title (0 - left, 0.5 - center, 1 - right)
         "x": 0.5,  # Adjust the x position of the title (0 - left, 0.5 - center, 1 - right)
        "y": 0.9,  # Adjust the y position of the title (0 - bottom, 1 - top)
        "yanchor": "top",  # Anchor point of the title (aligned to the top)
        # "font": {"size": 20},  # Font size of the title
        # "pad": {"t": 5}  
    },
    legend=dict(
        orientation="h",
    ),
    xaxis = dict (
        tickmode='linear',
        tick0 = min(df['survey_year']),
        dtick=1
    ),
    yaxis = dict (
        range=[0,100],
        title="Percentage value (%)",  # Add this line
    ),
      margin=dict(
        t=100  # Increase the top margin (adjust the value as needed)
    ),
    hovermode = 'x',
      )
    
    if getImage:
        img_bytes = io.BytesIO()
        py.write_image(fig, img_bytes, format='png')
        img_bytes.seek(0)
        return img_bytes
    else:
        plot_div = plot(fig, output_type='div', include_plotlyjs=False)
        if (df.drop('survey_year', axis=1) == 0.0).all().all():
            return None
        return plot_div


def load_dashboard(dashboard_filters,dropdown,isAdmin=False, getImage = False):
        # print("This is it", dashboard_filters, dropdown)
        plot1 = school_survey(dashboard_filters, isAdmin, getImage=getImage)
        plot2 = core_experience(dashboard_filters, getImage=getImage)
        plot3 = component_level(dashboard_filters, getImage=getImage)
        plot4 = core_experience_yearly(dashboard_filters, getImage=getImage)
        plot5 = implement_unified_sport_activity(dashboard_filters, getImage=getImage)
        plot6 = implement_youth_leadership_activity(dashboard_filters, getImage=getImage)
        plot7 = implement_school_engagement_activity(dashboard_filters, getImage=getImage)
        plot8 = sona_resources_useful(dashboard_filters, getImage=getImage)
        plot9 = school_survey_over_time(dashboard_filters, getImage=getImage)
        plot10 = core_experience_components(dashboard_filters, getImage=getImage, year = dashboard_filters['survey_taken_year'])


        context={
            # 'plot1':school_survey(dashboard_filters),
            # 'plot2':core_experience(dashboard_filters),
            # 'plot3':implementation_level(dashboard_filters),
            # 'plot4':core_experience_yearly(dashboard_filters),
            # 'plot5':implement_unified_sport_activity(dashboard_filters),
            # 'plot6':implement_youth_leadership_activity(dashboard_filters),
            # 'plot7':implement_school_engagement_activity(dashboard_filters),
            # 'plot8':sona_resources_useful(dashboard_filters),
            'form':dropdown,
            'plot_titles': {
            'plot1': "State Program response rate",
            'plot2': "Core Experience implementation",
            'plot3': "Components over time",
            'plot4': "Percentage of Core experience implementation over time",
            'plot5': "Percentage of schools implementing each Unified Sports activity",
            'plot6': "Percentage of schools implementing each Inclusive Youth Leadership activity",
            'plot7': "Percentage of schools implementing each Whole School Engagement activity",
            'plot8': "Percentage of liaisons who found SONA resources useful",
            'plot9': "Survey response rate over time",
            'plot10': "Components detail implementation",


        }
        }
        
        plots = [plot1, plot2, plot3, plot4, plot5, plot6, plot7, plot8, plot9, plot10]
        for i, plot in enumerate(plots, start=1):
            if plot is not None:
                context[f'plot{i}'] = plot
        return context



@login_required(login_url='../auth/login/')
def index(request):
    state = CustomUser.objects.values('state').filter(username=request.user)[0]
    state=state.get('state','None')
    admin = False
    context = {}
    if request.method=='GET':
        filter_state = state
        if state=='all':
            filter_state = 'AK'# on inital load some data has to be displayed so defaulting to ma
            admin = True
        
        context = {
            # 'plot1':school_survey(dashboard_filters),
            # 'plot2':core_experience(dashboard_filters),
            # 'plot3':implementation_level(dashboard_filters),
            # 'plot4':core_experience_yearly(dashboard_filters),
            # 'plot5':implement_unified_sport_activity(dashboard_filters),
            # 'plot6':implement_youth_leadership_activity(dashboard_filters),
            # 'plot7':implement_school_engagement_activity(dashboard_filters),
            # 'plot8':sona_resources_useful(dashboard_filters),
            'form':Filters(state=state_choices(state)),
            'plot_titles': {
            'plot1': "State Program response rate",
            'plot2': "Core Experience implementation",
            'plot3': "Components over time",
            'plot4': "Percentage of Core experience implementation over time",
            'plot5': "Percentage of schools implementing each Unified Sports activity",
            'plot6': "Percentage of schools implementing each Inclusive Youth Leadership activity",
            'plot7': "Percentage of schools implementing each Whole School Engagement activity",
            'plot8': "Percentage of liaisons who found SONA resources useful",
            'plot9': "Survey response rate over time",
            'plot10': "Components detail implementation",


            }
            }
    
    if request.method=='POST':
        # print(request)
        getState = CustomUser.objects.values('state').filter(username=request.user)[0]
        getState=getState.get('state','None')
        admin = False
        if getState == 'all':
            admin = True

        state = state_choices(state)
        post_data = request.POST.copy()  # Make a mutable copy of the POST data
        # if not isinstance(post_data.get('school_county'), str):
        dropdown = Filters(state,post_data)
        if not dropdown.is_valid():
            post_data['school_county'] = 'all'
            dropdown = Filters(state,post_data)
        


        # if not isinstance(dropdown.cleaned_data.get('school_county'), str):
        #         dropdown.cleaned_data['school_county'] = 'all'
        # print("===================",dropdown)
        # print("==State==", dropdown.is_valid())
        if not dropdown.is_valid():
            print(dropdown.errors)

        if dropdown.is_valid():
            #print('heloooooo')
            dashboard_filters = dropdown.cleaned_data
            # print("==State==", dashboard_filters)
            # print("========DROPDOWN========", dropdown)

            # print("@@@@@@@@@@@@", dashboard_filters)
            context = {
            # 'plot1':school_survey(dashboard_filters),
            # 'plot2':core_experience(dashboard_filters),
            # 'plot3':implementation_level(dashboard_filters),
            # 'plot4':core_experience_yearly(dashboard_filters),
            # 'plot5':implement_unified_sport_activity(dashboard_filters),
            # 'plot6':implement_youth_leadership_activity(dashboard_filters),
            # 'plot7':implement_school_engagement_activity(dashboard_filters),
            # 'plot8':sona_resources_useful(dashboard_filters),
            'form':dropdown,
            'plot_titles': {
            'plot1': "State Program response rate",
            'plot2': "Core Experience implementation",
            'plot3': "Components over time",
            'plot4': "Percentage of Core experience implementation over time",
            'plot5': "Percentage of schools implementing each Unified Sports activity",
            'plot6': "Percentage of schools implementing each Inclusive Youth Leadership activity",
            'plot7': "Percentage of schools implementing each Whole School Engagement activity",
            'plot8': "Percentage of liaisons who found SONA resources useful",
            'plot9': "Survey response rate over time",
            'plot10': "Components detail implementation",

            }
            }
            # context = load_dashboard(dashboard_filters,dropdown, isAdmin=admin)

                        # Create a new PDF in memory
            
    # cache.set('context', context, 300)  # Cache the context for 300 seconds

    return render(request, 'analytics/index_graph.html', context) 
         

def get_graph(request):
    # print("REQUEST CHECK", request)
    state = CustomUser.objects.values('state').filter(username=request.user)[0]
    state=state.get('state','None')
    # print(state)
    dashboard_filters = request.GET.copy()  # Get the filters from the request
    type = dashboard_filters['type']
    graph_no = dashboard_filters['graph_no']
    admin = False
    # print(dashboard_filters)
    dashboard_filters.pop('type')
    dashboard_filters.pop('graph_no')
    if state=='all':
        admin = True

    print("GRAPH NO>", graph_no)


    # print("NEW FUNCTION:", dashboard_filters, type, graph_no)
    if graph_no == '9':
        plot_div = school_survey_over_time(dashboard_filters)
    elif graph_no == '1':
        # print("NEW FUNCTION2:", graph_no)
        match type:
            case 'locale':
                plot_div = school_survey_locale_level(dashboard_filters, isAdmin = admin)
            case 'imp_level':
                plot_div = school_survey_implementation_level(dashboard_filters, isAdmin = admin)
            case 'school_level':
                plot_div = school_survey_school_level(dashboard_filters, isAdmin = admin)
            case 'survey_taken_year':
                plot_div = school_survey_year(dashboard_filters, isAdmin = admin)
            case 'reset':
                plot_div = school_survey(dashboard_filters, isAdmin = admin)
            case '':
                plot_div = school_survey(dashboard_filters, isAdmin = admin)
    elif graph_no == '2':
          plot_div = core_experience(dashboard_filters)
    elif graph_no == '3':
          plot_div = component_level(dashboard_filters)
    elif graph_no == '4':
          plot_div = core_experience_yearly(dashboard_filters)
    elif graph_no == '5':
            if type == 'reset':
                plot_div = implement_unified_sport_activity(dashboard_filters=dashboard_filters, type="default")

            else:    
                plot_div = implement_unified_sport_activity(dashboard_filters=dashboard_filters, type=type)
    elif graph_no == '6':
            if type == 'reset':
                plot_div = implement_youth_leadership_activity(dashboard_filters=dashboard_filters, type="default")

            else:    
                plot_div = implement_youth_leadership_activity(dashboard_filters=dashboard_filters, type=type)

    elif graph_no == '7':
            if type == 'reset':
                plot_div = implement_school_engagement_activity(dashboard_filters=dashboard_filters, type="default")

            else:    
                plot_div = implement_school_engagement_activity(dashboard_filters=dashboard_filters, type=type)
    elif graph_no == '8':
            if type == 'reset':
                plot_div = sona_resources_useful(dashboard_filters=dashboard_filters, type='default')
            else:
                plot_div = sona_resources_useful(dashboard_filters=dashboard_filters, type=type)
    elif graph_no == '10':
            print("GRAPH NO>", graph_no)
            plot_div = core_experience_components(dashboard_filters=dashboard_filters, year = dashboard_filters['survey_taken_year'])

 
    
    # plot_div = school_survey_implementation_level(dashboard_filters)
    # print("YOO NEW DATA", plot_div)
    return JsonResponse({'plot_div': mark_safe(plot_div)})
     

def get_counties(request):
    state_abv = request.GET.get('state_abv')
    counties = SchoolDetails.objects.filter(state_abv=state_abv).values_list('school_county', flat=True)
    county_dict = {}
    county_dict.update({county: county for county in counties if county != '-99'})
    county_dict.update({'all': 'All'} ) 

    # print(county_dict)
    return JsonResponse(county_dict, safe=False)
import base64
from reportlab.lib.utils import ImageReader  # Import ImageReader
from docx import Document
from report_generator.helpers import state_full_name
import os
from django.conf import settings
from django.http import HttpResponse
from docx.shared import Inches
from openai import OpenAI
import requests
import json

api_key = "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"
STATE='ma'
width_=6*914400 #image width in inches
height_=4*914400#image height in inches
def receive_graph_images(request):
    template_path = os.path.join(settings.BASE_DIR, 'static', 'generate_report_template.docx')
    doc = Document(template_path)
    # print("THE REQUEST", request)
    # data = request.POST.getlist  # This is key because you're sending JSON data
    # print("THE REQUEST DATA", data)
    # graphs = data  # This contains all your graph data
    data = json.loads(request.body)
    # print("THE DATA", data['graphs'])

    graphs = data['graphs']  # Access the graphs data

    # live_data = {
    #     '<<state>>': 'MA',
    #     '<<Figure1>>': 'image',
    #     '<<Figure2>>': 'image',
    #     '<<Figure3>>': 'image',
    #     # '<<Figure4>>': 'image',
    #     # '<<Figure5>>': 'image',
    #     # '<<Figure6>>': 'image',
    #     # '<<Figure7>>': 'image',
    #     # '<<Figure8>>': 'image',

    # }

    headers = {
    "Content-Type": "application/json",
    "Authorization": "Bearer sk-proj-PzPinLj1gmof4aDedOegT3BlbkFJRxBrDjAug2o3EU1zCT15"
    }
    arr = ["Implementation rates were consistent over the last two school years. There was a small increase in schools that implemented at least one Unified Sports activity (87% in 2022-2023, compared to 84% in 2021-2022). Similarly, there was a small increase in schools that implemented at least one Whole School Engagement activity (85% in 2022-2023, compared to 82% in 2021-2022). There was a larger increase in schools that implemented at least one Inclusive Youth Leadership activity (71% in 2022-2023, compared to 62% in 2021-2022). The initial recovery from COVID-19 observed in UCS schools in 2021-2022 appears to have continued in the 2022-2023 school year. See Figure 3 for the annual implementation rates of each core experience over the past ten years.", 
           "Starting in 2014-2015, Special Olympics (SO) categorized UCS implementation into three levels: • Full-implementation Unified Champion Schools implement at least one activity from all three core experiences. • Developing Unified Schools implement activities from Unified Sports and one other core experience. • Emerging Unified Schools implement activities from both Inclusive Youth Leadership and Whole School Engagement, or from just one of the three core experiences. UCS activities are most impactful when all core experiences are fully integrated and can work in tandem (Siperstein et al., 2019; Siperstein et al., 2017). Because of this, schools are encouraged to strive for Full-implementation status, with activities from all three core experiences. However, schools can still choose other combinations of the core experiences to cater to their unique contexts and needs. Among the 5,084 liaisons surveyed in 2022-2023, 60% were from Full-implementation schools, 21% were from Developing schools, and 16% were from Emerging schools. Historically, Full-implementation schools have been most common in the Liaison Survey sample each year. Compared to last year, there was a 9% increase in the percentage of Full-implementation schools, and a 4% and 8% decrease in Developing and Emerging schools. More schools that have had UCS for more than one year were identified as Full-implementation schools (65%) compared to new schools in their first year of UCS implementation (56%). Taking these findings together, the increase in Full-implementation schools is better accounted for by returning UCS schools that have either maintained their status over time or successfully transitioned to Full-implementation this year from Developing or Emerging last year.", 
           "The annual UCS Liaison Survey is pivotal to understanding UCS programming across schools, and liaisons have become an important source for assessing the program scope and impact nationwide. UCS liaisons are school officials designated as the point of contact between Special Olympics and their school, as well as leaders for UCS programming. Collaborating with state Special Olympics Programs, the CSDE evaluation team once again asked liaisons to share their insights and feedback as part of the annual evaluation. This year’s UCS Liaison Survey included a combination of closed-ended and open-ended questions that aimed to elicit rich details of UCS program implementation and its impact on schools and communities. The survey was divided into eight categories: a) liaison demographics and school characteristics, b) implementation of core experiences and activities, c) implementation support, d) Special Olympics’ resource awareness and usefulness, e) Special Olympics’ State Program support, f) funding, g) program sustainability, and h) impact of UCS programming on students and the school environment. This section of the annual evaluation is separated into multiple subsections. The methods subsection describes the processes involved in collecting data for this year’s Liaison Survey. Next, the following subsection describes UCS implementation across schools in 2022-2023. After that, additional analyses are presented regarding program sustainability and state-level or school-level implementation support. Lastly, this section concludes with a discussion of the impact of UCS on students and a school environment. Methods Between April and June 2023, the evaluation team contacted 7,350 liaisons across 51 State Programs. CSDE received responses from a total of 5,084 school liaisons, which was a national response rate of 69%. This is the highest response rate since the start of the COVID-19 pandemic (see Figure 1 for response rates since the 2015-2016 school year). This also represents an increase of 842 responses from the 2021-2022 school year. See Appendix B for a full breakdown of school responses by State Program.", 
           "Implementation rates were consistent over the last two school years. There was a small increase in schools that implemented at least one Unified Sports activity (87% in 2022-2023, compared to 84% in 2021-2022). Similarly, there was a small increase in schools that implemented at least one Whole School Engagement activity (85% in 2022-2023, compared to 82% in 2021-2022). There was a larger increase in schools that implemented at least one Inclusive Youth Leadership activity (71% in 2022-2023, compared to 62% in 2021-2022). The initial recovery from COVID-19 observed in UCS schools in 2021-2022 appears to have continued in the 2022-2023 school year. See Figure 3 for the annual implementation rates of each core experience over the past ten years.", "Starting in 2014-2015, Special Olympics (SO) categorized UCS implementation into three levels: • Full-implementation Unified Champion Schools implement at least one activity from all three core experiences. • Developing Unified Schools implement activities from Unified Sports and one other core experience. • Emerging Unified Schools implement activities from both Inclusive Youth Leadership and Whole School Engagement, or from just one of the three core experiences. UCS activities are most impactful when all core experiences are fully integrated and can work in tandem (Siperstein et al., 2019; Siperstein et al., 2017). Because of this, schools are encouraged to strive for Full-implementation status, with activities from all three core experiences. However, schools can still choose other combinations of the core experiences to cater to their unique contexts and needs. Among the 5,084 liaisons surveyed in 2022-2023, 60% were from Full-implementation schools, 21% were from Developing schools, and 16% were from Emerging schools. Historically, Full-implementation schools have been most common in the Liaison Survey sample each year. Compared to last year, there was a 9% increase in the percentage of Full-implementation schools, and a 4% and 8% decrease in Developing and Emerging schools. More schools that have had UCS for more than one year were identified as Full-implementation schools (65%) compared to new schools in their first year of UCS implementation (56%). Taking these findings together, the increase in Full-implementation schools is better accounted for by returning UCS schools that have either maintained their status over time or successfully transitioned to Full-implementation this year from Developing or Emerging last year.", 
           "The annual UCS Liaison Survey is pivotal to understanding UCS programming across schools, and liaisons have become an important source for assessing the program scope and impact nationwide. UCS liaisons are school officials designated as the point of contact between Special Olympics and their school, as well as leaders for UCS programming. Collaborating with state Special Olympics Programs, the CSDE evaluation team once again asked liaisons to share their insights and feedback as part of the annual evaluation. This year’s UCS Liaison Survey included a combination of closed-ended and open-ended questions that aimed to elicit rich details of UCS program implementation and its impact on schools and communities. The survey was divided into eight categories: a) liaison demographics and school characteristics, b) implementation of core experiences and activities, c) implementation support, d) Special Olympics’ resource awareness and usefulness, e) Special Olympics’ State Program support, f) funding, g) program sustainability, and h) impact of UCS programming on students and the school environment. This section of the annual evaluation is separated into multiple subsections. The methods subsection describes the processes involved in collecting data for this year’s Liaison Survey. Next, the following subsection describes UCS implementation across schools in 2022-2023. After that, additional analyses are presented regarding program sustainability and state-level or school-level implementation support. Lastly, this section concludes with a discussion of the impact of UCS on students and a school environment. Methods Between April and June 2023, the evaluation team contacted 7,350 liaisons across 51 State Programs. CSDE received responses from a total of 5,084 school liaisons, which was a national response rate of 69%. This is the highest response rate since the start of the COVID-19 pandemic (see Figure 1 for response rates since the 2015-2016 school year). This also represents an increase of 842 responses from the 2021-2022 school year. See Appendix B for a full breakdown of school responses by State Program.", 
           "Inclusive Youth Leadership activities empower students to be leaders and to develop social skills such as advocacy and decision-making. A primary goal of Inclusive Youth Leadership is to offer students with and without IDD opportunities to nurture these skills, share their unique experiences, and enact changes in their communities. In 2022-2023, schools implemented an average of one to two Inclusive Youth Leadership activities. Schools’ implementation of Inclusive Youth Leadership activities was similar to 2021-2022 and pre-pandemic years. Unified Club (72%) continued to be the most frequently implemented activity nationally, followed by Inclusive Leadership Training/Class (39%). Consistent with previous years, the implementation of most Inclusive Youth Leadership activities varied by school level. As expected, Young Athletes Volunteers were mostly offered within elementary schools, while Unified Club, SO Youth Summit, and SO Youth Activation Committees were more common at the high school level. Similar to 2021-22, liaisons reported on the level of student participation for each Inclusive Youth Leadership activity. The proportion of schools that had inclusive participation (e.g., including students with and without IDD) for each activity ranged from 83% (for Young Athletes Volunteers) to 95% (for Unified Club). Some liaisons (n = 372) disclosed the reasons why only students with IDD or only students without IDD participated in inclusive activities, including a) the nature of the activity, b) the need for more time and support, c) scheduling conflicts, d) limitations in student group participation, and e) issues with transportation. A common reason for why only one student group participated was because of the activity design (n = 74). For example, some schools utilized Inclusive Youth Leadership Training/Class or Youth Athlete Volunteers as an opportunity for students without IDD to learn more about disabilities, while inclusion alongside students with IDD was promoted in other activities and events. Another main reason related to a lack of time and support (n = 70). Some liaisons reported that because their school was new to UCS or transitioning with new staff, it became challenging to implement inclusive activities this year. Other liaisons (n = 59) noted scheduling conflicts. Since students with and without IDD often had different schedules during the school day, it was difficult to have all of them participate in the same activity. Additionally, some schools had difficulty attracting diverse student participation into these activities (n = 52), particularly if they only served students with disabilities or had very few students with IDD enrolled this year. Lastly, a small group of liaisons (n = 15) mentioned that due to safety concerns, transportation became an issue when inviting students with IDD to participate in certain activities. A final aspect of exploring this year’s Inclusive Youth Leadership activities focused on a deeper analysis of a school’s Unified Club. As the most frequently implemented Inclusive Youth Leadership activity for years, Unified Club offers a school-based hub to gather students with and without IDD together. Prior research and evaluation findings showed that students who are actively involved in Unified Club gain more positive experiences related to taking lead roles and developing a sense of responsibility. In the 2022-2023 evaluation, 31% of schools had a Unified Club that met at least once per week, and 56% of schools had a club that met monthly. When club members met, they focused on social emotional learning skills (68%), leadership (60%), event planning (46%), and advocacy skills for students with and without IDD (46%). In contrast, far fewer schools used the Unified Club to promote college and career readiness skills (9%).",
             "Unified Sports is an essential component of UCS programming. Unified Sports activities are designed to create opportunities for students with and without IDD to train, compete, and develop understanding and friendship together. On average, schools implemented two Unified Sports activities in the 2022-2023 school year. The two most common activities implemented across schools were Unified Sports teams (65%) and Unified PE (63%). The implementation of each Unified Sports activity remained consistent with last year’s evaluation. In the 2020-2021 school year, 39% of liaisons reported having a Unified Sports team. This number increased to 60% in the 2021-22 school year and 65% in the 2022-2023 school year. Looking back, although the proportion of schools that have Unified Sports teams has not fully returned to pre-pandemic levels, the steadily increasing implementation rate continues to demonstrate ongoing recovery of UCS programming since the start of the COVID-19 pandemic. This year, urban schools had somewhat higher rates of Unified PE and Unified Fitness compared to other locales. Conversely, Unified Sports was somewhat more prevalent in suburban, town, and rural schools compared to urban schools. Other Unified Sports activities were similarly implemented across school locales. Overall, these findings suggest that there is a small difference in the Unified Sports activities that a school chooses to implement based on whether they are in an urban or nonurban area. Lastly, liaisons reported the makeup of their Unified Sports teams as well as the activities that Unified Sports teams had during the school year. Unified Sports teams are designed to bring students with and without IDD together for various sports in both competitive and recreational models. In 2022-23, 60% of schools that had at least one Unified Sports activity offered a Unified Sports team for two or more seasons. Of schools that had a Unified Sports team, 90% of schools had competition against Unified Sports teams from another school. The implementation of Unified Sports teams for multiple seasons, especially in the competitive model, was more common among middle and high schools. On average, schools with a Unified Sports team had two to three coaches this year and 64% of the coaches were trained or certified by Special Olympics. Across schools that had Unified Sports teams this year, nearly 70% of coaches were trained or certified by Special Olympics and 56% of coaches were certified through the National Federation of High Schools (NFHS). Most liaisons reported that their coaches were certified by both NFHS and Special Olympics (46%), 21% of liaisons reported that their coaches were only certified by Special Olympics, and 5% of liaisons reported that their coaches were only certified by NFHS.", 
             "Whole School Engagement provides opportunities for all students to engage in UCS activities, facilitating a culture of social inclusion. Because Whole School Engagement events and activities include the largest number of students in a school, they can raise awareness of the capabilities and contributions of students with IDD while promoting an inclusive school culture. In the 2022-23 school year, schools on average implemented between two to three Whole School Engagement activities. Overall, the percentage of schools implementing each Whole School Engagement activity is consistent with the findings from last year’s annual evaluation. There was an increase in the percentage of schools that held a Fans in the Stands/Unified Sports Pep Rally (e.g., 50% of schools implemented this event in 2022-2023 compared to 41% of schools in 2021-2022). The consistency of Whole School Engagement activities over the last two school years supports the continued recovery from COVID-19 that was observed in last year’s annual evaluation. Looking at activity implementation by locale, urban schools showed similar rates of implementation for each activity compared to town, rural, and suburban locales, except for Fans in the Stands/Unified Sports Pep Rally and fundraising. Only 32% of urban schools had at least one fundraising event or activity, compared to 43% to 50% of suburban, town, or rural schools. Similarly, only 44% of urban schools had a Fans in the Stands/Unified Sports Pep Rally, compared to 49% to 52% of suburban, town, or rural schools. These differences are small, but they may reflect differences in how urban schools implement Whole School Engagement activities.", 
             ""]

    placeholders = [f"<<Figure{i}>>" for i in range(1, 11)]  # Adjust the range if more placeholders
    # print(placeholders)
    for index, (graph, placeholder) in enumerate(zip(graphs, placeholders)):
            
            found_image_placeholder = False
            for paragraph in doc.paragraphs:
                    if placeholder in paragraph.text:
                            # print("Analyze and give the report of the data in the graph from the data given:",  image_data,". Give the result in such a way that it can be direclty used in a report file. Dont mention these legends data:", disabled_legends,". keep the description professional, include percentage values too. ")


                            image_data_img = graph['image_data_img']
                            disabled_legends = graph['disabled_legends']
                            image_data = graph['image_data']
                            paragraph.text = paragraph.text.replace(placeholder, '')

                            if image_data_img:
                                format, imgstr = image_data_img.split(';base64,')
                                image_bytes = base64.b64decode(imgstr)
                                image_stream = BytesIO(image_bytes)

                                # Debug: Print the size of the image stream
                                # print("Image stream size:", len(image_bytes))

                                run = paragraph.add_run()
                                # run.text = run.text.replace(placeholder, '')

                                # image_stream.seek(0)  # Ensure the stream is at the start
                                run.add_picture(image_stream, width=Inches(7), height=Inches(4))
                                payload = {
                                    "model": "gpt-4o",
                                    "messages": [
                                        {
                                            "role": "user",
                                            "content": f"Analyze and give the summary report of the data in the graph from the data given:  {image_data}. Give the result in such a way that it can be direclty used in a docx report file. Ignore these name from the data: {disabled_legends}. keep the description professional, use this template: {arr[index]}"
                                        }
                                    ],
                                    "max_tokens": 2000
                                    }
                                print("Image added to document.", payload)
                                response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload)
                                print(response.json())
                                content = response.json()['choices'][0]['message']['content']
                                text_run = paragraph.add_run()
                                text_run.text = content

                                
                                found_image_placeholder = True
                    
    # After processing, remove any remaining placeholders
    for placeholder in placeholders:
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, '')

    if found_image_placeholder:
        output_path = os.path.join(settings.BASE_DIR, 'static', 'test.docx')
        doc.save(output_path)
        print("Document saved to:", output_path)
        with open(output_path, 'rb') as docx_file:
            response = HttpResponse(docx_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename="graph_analysis_report.docx"'
            return response
    else:
        print("No image placeholder found in document.")

        return HttpResponse("No placeholders were processed.", content_type='text/plain')

        

    # If not POST or other error, return error response
    return JsonResponse({'status': 'error'}, status=400)

'''
Utility functions
'''

'''
input list or single dict 
return list of percent values or returns nested dict with value and percent
op sample :'state': {None: {'value': 0, 'percent_val': 0.0}, 'Yes': {'value': 4, 'percent_val': 7.0}, 'No': {'value': 53, 'percent_val': 93.0}}}
'''
def percentage_values(total_values, column_name = "", compare_type = "default"):
    # print("percentage_values", column_name, total_values, compare_type)
    if type(total_values) == list:
        values = total_values
        return [round(((i/sum(values))*100),2) for i in values]
    elif type(total_values) == dict:
        data = list(total_values.values())
        try:
            percent_arr=[round(((i/sum(data))*100),1) for i in data ] 
        except ZeroDivisionError:
            percent_arr = [0]*len(data)
        res = {key:{'value':total_values[key],'percent_val':percent_arr[i]} for i,key in enumerate(total_values.keys()) }
        return res
    elif isinstance(total_values, QuerySet):
        result = {}
        sum_of_totals = sum(item['total'] for item in total_values)
        for item in total_values:
            column_name_val = item[column_name]
            if column_name_val not in result:
                result[column_name_val] = []
            percentage = round((item['total'] / sum_of_totals) * 100, 2) if sum_of_totals != 0 else 0
            match compare_type:
                case 'imp_level':
                    result[column_name_val].append({'implementation_level': item['implementation_level'], 'total': item['total'], 'percentage': percentage})
                case 'locale':
                    result[column_name_val].append({'grouped_locale': item['grouped_locale'], 'total': item['total'], 'percentage': percentage})
                case 'school_level':
                    result[column_name_val].append({'gradeLevel_WithPreschool': item['gradeLevel_WithPreschool'], 'total': item['total'], 'percentage': percentage})
                case 'survey_taken_year':
                    # print(result)
                    result[column_name_val].append({'survey_taken_year': item['survey_taken_year'], 'total': item['total'], 'percentage': percentage})
                    # print(result[column_name_val])
                    # print("WORKING!!")

        # print("the new result", result)
        return result
    
def state_choices(state):#used for drop down in filters
    STATE_CHOICES = []
    # print(state)
    STATE_CHOICES_RAW= list(SchoolDetails.objects.values_list('state_abv','school_state').distinct())
    # print(STATE_CHOICES_RAW)

    if state =='all' or state== STATE_CHOICES_RAW:
        for val in STATE_CHOICES_RAW:
            if val[0]!='-99' and None not in val:
                # print(val)
                STATE_CHOICES.append(val)
                STATE_CHOICES.sort()
        return STATE_CHOICES
    else:
        for val in STATE_CHOICES_RAW:
            if val[1]==state:
                return [(val[0],val[1])]
    return None
'''
query function for below six graphs
as pattern is same only column names are changing i used this method for fetching data
#none/null values are by default counted as o in this query
op sample:# {'sports_sports_teams': {'national': {'No': 2094, None: 0, 'Yes': 2033}, 'state': {None: 0, 'No': 24, 'Yes': 33}}, 'sports_unified_PE': {'national': {'No': 1945, None: 0, 'Yes': 2182}, 'state': {None: 0, 'No': 22, 'Yes': 35}}, 'sports_unified_fitness': {'national': {'No': 3292, None: 0, 'Yes': 835}, 'state': {None: 0, 'No': 44, 'Yes': 13}}, 'sports_unified_esports': {'national': {'No': 3903, None: 0, 'Yes': 224}, 'state': {None: 0, 'Yes': 4, 'No': 53}}, 'sports_young_athletes': {'national': {'No': 3450, None: 0, 'Yes': 677}, 'state': {None: 0, 'No': 46, 'Yes': 11}}, 'sports_unified_developmental_sports': {'national': {'No': 3505, None: 0, 'Yes': 622}, 'state': {None: 0, 'No': 49, 'Yes': 8}}}

'''
def main_query(column_name,filters,key,compareType="default"): 
    # print("FILTERS NEW:", filters, column_name)
    if key=='state':
        filters=filters
    if key=='all':
        filters=filters.copy()
        if 'state_abv' in filters:
            filters.pop('state_abv')
        if 'state_abv__in' in filters:
            filters.pop('state_abv__in')
        if 'county' in filters:
            filters.pop('county')

        
    # print("Main_QUery_filters", compareType, filters)
    # print("Filters wow", filters)

    data = {}
    match compareType:
        case "default":
            data = dict(SchoolDetails.objects.values_list(column_name).filter(**filters).annotate(total = Count(column_name)))
        case "imp_level":
            data = (SchoolDetails.objects.values(column_name, "implementation_level").filter(**filters).annotate(total = Count(column_name)))
        case "locale":
            data = (SchoolDetails.objects.values(column_name).filter(**filters).annotate(total = Count(column_name)))

            # data = dict(SchoolDetails.objects.values_list(column_name, "locale").filter(**filters).annotate(total = Count(column_name)))
            data = data.annotate(
             grouped_locale=Case(
            When(locale__startswith='City', then=Value('City')),
            When(locale__startswith='Rural', then=Value('Rural')),
            When(locale__startswith='Suburb', then=Value('Suburb')),
            When(locale__startswith='Town', then=Value('Town')),
            default=Value('Other'),
            output_field=CharField(),
            )
            )
            data = data.values(column_name, "grouped_locale").annotate(total = Count(column_name))

        case "school_level":
            data = (SchoolDetails.objects.values(column_name, "gradeLevel_WithPreschool").filter(**filters).annotate(total = Count(column_name)))
        case "survey_taken_year":

            filters.pop('survey_taken_year',None)

    
    # print(filters)
    
    #  SchoolDetails.objects.values('school_state','survey_taken', 'survey_taken_year').filter(**filters).exclude(school_state='-99').annotate(total = Count('survey_taken')).order_by('school_state')
            data = (SchoolDetails.objects.values(column_name, "survey_taken_year").filter(**filters).annotate(total = Count(column_name)))
            # print("RESULT DATA:", data)
    
    # print("MAIN QUERY PRINT", compareType, data)
    #test query to run in terminal: dict(SchoolDetails.objects.values_list('sports_sports_teams').filter(state_abv='sca',survey_taken_year=2022).annotate(total = Count('sports_sports_teams')))
    return data




'''
LOGIC
'''
def implement_unified_sport_activity(dashboard_filters, type="default", getImage = False):
    response={'sports_sports_teams':{}, 'sports_unified_PE':{},'sports_unified_fitness':{},
             'sports_unified_esports':{},'sports_young_athletes':{},'sports_unified_developmental_sports':{}}
    filters=filter_set(dashboard_filters)
    # print("implement_unified_sport_activity filters", filters)
    filters = add_in_forFilters(filters)

    comparedBy = ""
    match type:
        case 'imp_level':
            comparedBy = "compared by Implementation Level"
        case 'locale':
            comparedBy = "compared by Locale"
        case 'school_level':
            comparedBy = "compared by School Level"
        case 'survey_taken_year':
            comparedBy = "compared by Survey Year" 
    for column_name in response.keys():
        response[column_name]['national']=percentage_values(main_query(column_name,filters,key='all')) 
        response[column_name]['state']=percentage_values(main_query(column_name,filters,key='state', compareType=type), column_name, compare_type=type)
    y_axis = ['Unified Sports Teams', 'Unified PE', 'Unified fitness','Unified esports', 'Young athletes(participants)', 'Unified Developmental Sports']
    title='Percentage of schools implementing each <br> Unified Sports activity'.format(dashboard_filters['state_abv'], comparedBy)
    state=dashboard_filters['state_abv']#adding state to the response for graph lables
    emptyGraph = True
    # print("implement_unified_sport_activity state", response)

    for column_name in response.keys():
        if response[column_name]['state']:
            # print("Found state", response[column_name]['state'])
            emptyGraph = False
    # print("SOMETHING!!!!", emptyGraph)
    if emptyGraph:
        return None
    else:
        return horizontal_bar_graph(response,y_axis,title,state,compare_type=type,getImage=getImage)
    
    


def implement_youth_leadership_activity(dashboard_filters, type="default", getImage = False):
    response = {'leadership_unified_inclusive_club':{},'leadership_youth_training':{},'leadership_athletes_volunteer':{},
               'leadership_youth_summit':{},'leadership_activation_committe':{}}
    filters=filter_set(dashboard_filters)
    filters = add_in_forFilters(filters)
    # print("===========THE TYPE:", type)
    comparedBy = ""
    match type:
        case 'imp_level':
            comparedBy = "compared by Implementation Level"
        case 'locale':
            comparedBy = "compared by Locale"
        case 'school_level':
            comparedBy = "compared by School Level"
        case 'survey_taken_year':
            comparedBy = "compared by Survey Year"   
    
    for column_name in response.keys():
        response[column_name]['national']=percentage_values(main_query(column_name,filters,key='all')) 
        response[column_name]['state']=percentage_values(main_query(column_name,filters,key='state', compareType=type), column_name, compare_type=type)
    y_axis=['Unifed/Inclusive Club','Inclusive Youth Leadership Training/Class','Young Athletes Volunteers','Youth summit','Youth Activation Committee']
    title='Percentage of schools implementing each <br>  Youth Leadership activity'.format(dashboard_filters['state_abv'], comparedBy)
    state=dashboard_filters['state_abv']#adding state to the response for graph lables
    
    
    emptyGraph = True
    # print("implement_unified_sport_activity state", response)

    for column_name in response.keys():
        if response[column_name]['state']:
            # print("Found state", response[column_name]['state'])
            emptyGraph = False

    if emptyGraph:
        return None
    else:
        return horizontal_bar_graph(response,y_axis,title,state,compare_type=type, getImage= getImage)
    
    # return horizontal_bar_graph(response,y_axis,title,state,compare_type=type)

    


def implement_school_engagement_activity(dashboard_filters, type="default", getImage = False):
    response = {'engagement_spread_word_campaign':{},'engagement_fansinstands':{},'engagement_sports_day':{},
                'engagement_fundraisingevent':{},'engagement_SO_play':{},'engagement_fitness_challenge':{}}
    
    filters=filter_set(dashboard_filters)
    filters = add_in_forFilters(filters)
    comparedBy = ""
    match type:
        case 'imp_level':
            comparedBy = "compared by Implementation Level"
        case 'locale':
            comparedBy = "compared by Locale"
        case 'school_level':
            comparedBy = "compared by School Level"
        case 'survey_taken_year':
            comparedBy = "compared by Survey Year" 
    for column_name in response.keys():
        response[column_name]['national']=percentage_values(main_query(column_name,filters,key='all',)) 
        response[column_name]['state']=percentage_values(main_query(column_name,filters,key='state',compareType=type), column_name, compare_type=type)
    y_axis=['Spread the word'+'<br>'+'Inclusion/Respect/Disability' +'<br>'+'Awareness Campaign','Unified Sports pep Rally','Unified Sports Day/Festival','Fundraising events /activities','Special Olympics play/performance','Unified Fitness Challenge']
    # print("STATE ABV CHECK", dashboard_filters['state_abv'])
    title='Percentage of schools implementing each <br> Inclusive Whole School Engagement activity'.format(dashboard_filters['state_abv'], comparedBy)
    state=dashboard_filters['state_abv']#adding state to the response for graph lables
    emptyGraph = True
    # print("implement_unified_sport_activity state", response)

    for column_name in response.keys():
        if response[column_name]['state']:
            # print("Found state", response[column_name]['state'])
            emptyGraph = False

    if emptyGraph:
        return None
    else:
        return horizontal_bar_graph(response,y_axis,title,state,compare_type=type, getImage = getImage)
    
    # return horizontal_bar_graph(response,y_axis,title,state,compare_type=type)


def sona_resources_useful(dashboard_filters, type="default", getImage = False):
    response={'elementary_school_playbook':{},'middle_level_playbook':{},'high_school_playbook':{},'special_olympics_state_playbook':{},'special_olympics_fitness_guide_for_schools':{},'unified_physical_education_resource':{},
              'special_olympics_young_athletes_activity_guide':{},'inclusive_youth_leadership_training_faciliatator_guide':{},'planning_and_hosting_a_youth_leadership_experience':{},'unified_classoom_lessons_and_activities':{},'generation_unified_youtube_channel_or_videos':{},'inclusion_tiles_game_or_facilitator_guide':{}}
    column_names=response.keys()
    filters=filter_set(dashboard_filters)
    filters = add_in_forFilters(filters)
    # print("sona", filters)
    for column_name in response.keys():
        response[column_name]['national']=percentage_values(main_query(column_name,filters,key='all')) 
        response[column_name]['state']=percentage_values(main_query(column_name,filters,key='state',compareType=type), column_name, compare_type=type)
    # state=dashboard_filters['state_abv']#adding state to the response for graph lables
    ##seperating yes and no keys from each parent key
    # print("sona 1", response, type)
    state_no=[]
    nation_no=[]
    state_yes=[]
    nation_yes=[]
    state_yes_val=[]
    state_no_val=[]
    nation_yes_val=[]
    if type == "default":
        for val in column_names:
            n_keys=response[val]['national'].keys()
            s_keys=response[val]['state'].keys()
            for key in n_keys:
                if key=='1':
                    # nation_no.append(response[val]['national'].get(key,{}).get('percent_val',0))
                # elif key!=None and key!='0': 
                    #logic here is the response has only three keys for all columns so if its not these two then it should be the column name
                    #saves time to avoid writing down all column yes names and then verifying
                    nation_yes.append(response[val]['national'].get(key,{}).get('percent_val',0))
                    nation_yes_val.append(response[val]['national'].get(key,{}).get('value',0))

            for key in s_keys:
                if key=='1':
                    
                    # state_no.append(response[val]['state'].get(key,{}).get('percent_val',0))
                # elif key!=None and key!='0':
                    #logic here is the response has only three keys for all columns so if its not these two then it should be the column name
                    #saves time to avoid writing down all column yes names and then verifying
                    state_yes.append(response[val]['state'].get(key,{}).get('percent_val',0))
                    state_yes_val.append(response[val]['state'].get(key,{}).get('value',0))

    else:
        for val in column_names:
            n_keys=response[val]['national'].keys()
            s_keys=response[val]['state'].keys()
            # print("n keys", n_keys, "s keys", s_keys)
            for key in n_keys:
                if key=='1':
                    # nation_no.append(response[val]['national'].get(key,{}).get('percent_val',0))
                # elif key!=None and key!='0': 
                    nation_yes.append(response[val]['national'].get(key,{}).get('percent_val',0))

            for key in s_keys:
                if key=='1':
                    # state_no.append(sum(item.get('percentage', 0) for item in response[val]['state'].get(key, [])))
                # elif key!=None and key!='0':
                    state_yes.append([{'column': val, 'percentage': item.get('percentage', 0)} for item in response[val]['state'].get(key, [])])
                    #logic here is the response has only three keys for all columns so if its not these two then it should be the column name
                    #saves time to avoid writing down all column yes names and then verifying
                    # state_yes.append(response[val]['state'].get(key,{}).get('percent_val',0))
    new_response = {'state_yes':state_yes,'nation_yes':nation_yes, 'state_yes_val': state_yes_val, 'nation_yes_val': nation_yes_val}
    # print("SONA RESPONSE", new_response)
    y_axis=['Elementary School Playbook: A Guide for Grades K-5','Middle Level Playbook: A Guide for Grades 5-8','High School Playbook','Special Olympics State Playbook','Special Olympics Fitness Guide for Schools','Unified Physical Education Resource',
            'Special Olympics Young Athletes Activity Guide','Inclusive Youth Leadership Training: Faciliatator Guide','Planning and Hosting a Youth Leadership Experience','Unified Classoom lessons and activities','Generation Unified Youtube channel or videos',
            'Inclusion Tiles game or facilitator guide']
    
    title='Percentage of liaisons who found SONA resources useful'.format(dashboard_filters['state_abv'])
    state=dashboard_filters['state_abv']
    # return horizontal_bar_graph(response,y_axis,title,state, compare_type=type)
    # print(new_response)
    if new_response["state_yes_val"] or new_response["nation_yes_val"]:
        return horizontal_stacked_bar(new_response,y_axis,title,state, compare_type=type, getImage = getImage)
    else:
        return None


'''   
CHARTS
'''
def horizontal_bar_graph(response,y_axis,heading,state,compare_type="default", getImage = False):
    # print("horizontal_bar_graph", response, compare_type)

    fig = go.Figure()
    colors = ['rgba(200, 200, 220, 0.8)', 'rgba(168, 184, 134, 0.8)', 'rgba(170, 180, 200, 0.8)', 'rgba(109, 154, 168, 0.8)', 'rgba(120, 130, 160, 0.8)', 'rgba(140, 138, 173, 0.8)']     # Define your colors here with opacity
    # Get the list of dictionaries for key '1'
    # Create a bar for each dictionary in state_data
    if compare_type == "default":
        fig.add_trace(go.Bar(
        y=y_axis,
        x=[round(response[val]['national'].get('1',{}).get('percent_val',0)) for val in response if response],
        name='National',
        customdata=[response[val]['national'].get('1',{}).get('value',0) for val in response if response],
        orientation='h',
        hovertemplate='%{x}%, %{customdata} schools, %{y}',

        visible = "legendonly",
        marker=dict(
            color='rgba(246, 78, 139, 0.6)',
            line=dict(color='rgba(246, 78, 139, 1.0)', width=0)
            
                  )
            ))
        # print("fig 1")

        # print("PRINT DEFAULT")
        fig.add_trace(go.Bar(
        y=y_axis,
        x=[round(response[val]['state'].get('1',{}).get('percent_val',0)) for val in response if response ],
        customdata=[response[val]['state'].get('1',{}).get('value',0) for val in response if response],

        name=state,
        orientation='h',
        hovertemplate='%{x}%, %{customdata} schools, %{y}',

        marker=dict(
            color='rgba(58, 71, 80, 0.6)',
            line=dict(color='rgba(58, 71, 80, 1.0)', width=0)
        )
      ))
    else:
        for i, val in enumerate(response):
            # print("horizontal_graph_state_data1", val)
            state_data = response[val]['state']
            # Create a bar for each key in state_data
            # print("horizontal_graph_state_data", state_data)
            for key in state_data:
                if key is '1':  # Exclude None key
                    state_data[key] = sorted(state_data[key], key=lambda x: x.get('percentage', 0), reverse=True)
                    for j, level_data in enumerate(state_data[key]):
                        # print("horizontal_graph_level_data", level_data, y_axis[i], i)
                        fig.add_trace(go.Bar(
                            y=[y_axis[i]],
                            x=[round(level_data.get('percentage', 0))],
                            name=f"{state} {y_axis[i]} Level {get_horizontal_bar_legend_name(level_data, compare_type)} ",
                            orientation='h',
                            offsetgroup=1,
                            hovertemplate=f"{round(level_data.get('percentage', 0))}%, {level_data.get('total', 0)} schools, Level {get_horizontal_bar_legend_name(level_data, compare_type)}",
                            width=0.5,
                            # base=j*1.1,  # Shift each bar slightly
                            # Add this line to increase the width of the bars
                            marker=dict(
                                color=colors[j % len(colors)],  # Use the index to select a color
                                line=dict(color=colors[j % len(colors)], width=0)
                        )))
                        # print(f"fig {val} Level {level_data.get('implementation_level')}")


    # fig.update_traces(visible=False, showLegend=True selector=dict(name="National"))


    fig.update_layout( 
        title={
            'text': heading,
            'x': 0.5,  # Center the title
            # 'xanchor': 'center',  # Ensure the title doesn't move
        },
        barmode='group',
        hovermode= 'y',
       
        height=500,  # Adjust the height of the graph
        # margin=dict(l=100, r=50, b=100, t=100, pad=4),
        xaxis=dict(title="Percentage value (%)"),  # This line is added to reverse the order of categories on y-axis
         legend=dict(
             traceorder='reversed'
              )
        
    )
    if getImage:
        img_bytes = io.BytesIO()
        py.write_image(fig, img_bytes, format='png')
        img_bytes.seek(0)
        return img_bytes
    else:
        plot_div = plot(fig, output_type='div', include_plotlyjs=False)
        return plot_div
    
def get_horizontal_bar_legend_name(data, type="default"):
    match type:
        case 'imp_level':
            match data['implementation_level']:
                case '1':
                     return 'Emerging'
                case '2':
                     return 'Developing'
                case '3':
                     return 'Full Implementation'
        case 'locale':
            return data['grouped_locale']
        case 'school_level':
            match data['gradeLevel_WithPreschool']:
                case '1.00':
                    return 'Elementary'
                case '2.00':
                    return 'High'
                case '3.00':
                    return 'Middle'
                case '4.00':
                    return 'Preschool'
        case 'survey_taken_year':
            match data['survey_taken_year']:
                case 2018:
                    return '2018'
                case 2019:
                    return '2019'
                case 2020:
                    return '2020'
                case 2021:
                    return '2021'
                case 2022:
                    return '2022'
                case 2023:
                    return '2023'

                
    return 'none'            





def horizontal_stacked_bar(response,y_axis,heading,state, compare_type="default", getImage = False):
    # print("SONA RESULT", response)
    fig = go.Figure()
    fig.add_trace(go.Bar(
        y=y_axis,
        x=[round(val) for val in response['nation_yes']],
        customdata=response['nation_yes_val'],
        name='National',
        visible = "legendonly",
        hovertemplate='%{x}%, %{customdata} schools, %{y}',
        orientation='h',
        marker=dict(
            color='rgba(99, 110, 250, 0.8)',
            line=dict(color='rgba(99, 110, 250, 1.0)', width=0)
        )
    ))
    if compare_type == "default":
        fig.add_trace(go.Bar(
            y=y_axis,
            x=[round(val) for val in response['state_yes']],
            customdata=response['state_yes_val'],
            name=state,
            hovertemplate='%{x}%, %{customdata} schools, %{y}',
            orientation='h',
            marker=dict(
                color='rgba(139, 144, 209, 0.8)',
                line=dict(color='rgba(139, 144, 209, 1)', width=0)
            )
        ))
    else:
        for i, val in enumerate(response['state_yes']):
            # print("horizontal_stacked_bar", val)
            # state_data = response[val]['state']
            # # Create a bar for each key in state_data
            # print("horizontal_graph_state_data", state_data)
            # for key in state_data:
            #     if key is '1':  # Exclude None key
            val = sorted(val, key=lambda x: x.get('percentage', 0), reverse=True)
            # print("horizontal_stacked_bar val", val)
                    # for j, level_data in enumerate(state_data[key]):
                    #     print("horizontal_graph_level_data", level_data, y_axis[i], i)
                    #     fig.add_trace(go.Bar(
                    #         y=[y_axis[i]],
                    #         x=[level_data.get('percentage', 0)],
                    #         name=f"{state} {y_axis[i]} Level {get_horizontal_bar_legend_name(level_data, compare_type)} ",
                    #         orientation='h',
                    #         offsetgroup=1,
                    #         hovertemplate=f"{level_data.get('percentage', 0)} Level {get_horizontal_bar_legend_name(level_data, compare_type)}",
                    #         width=0.5,  # Add this line to increase the width of the bars
                    #         marker=dict(
                    #             color=colors[j % len(colors)],  # Use the index to select a color
                    #             line=dict(color=colors[j % len(colors)], width=0)
                    #     )))
                    #     print(f"fig {val} Level {level_data.get('implementation_level')}")

    fig.update_layout(title=heading,barmode='group',hovermode= 'y',xaxis_range=[0,100], xaxis_title="Percentage value (%)", width=1400, height=500)
    if getImage:
        img_bytes = io.BytesIO()
        py.write_image(fig, img_bytes, format='png')
        img_bytes.seek(0)
        return img_bytes   
    else:    
        plot_div = plot(fig, output_type='div', include_plotlyjs=False)
        return plot_div

