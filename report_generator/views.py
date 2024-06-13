from docx import Document
from analytics.views import implement_unified_sport_activity,implement_youth_leadership_activity,implement_school_engagement_activity,sona_resources_useful
from report_generator.helpers import state_full_name

STATE='ma'
width_=6*914400 #image width in inches
height_=4*914400#image height in inches
dashboard_filters={'state_abv':STATE,'survey_taken_year':2022}



def fill_placeholders():
    # Open the existing template
    doc = Document('../static/prototype_week1.docx')
    live_data = {
        '<<state>>': state_full_name(STATE),
        'placeholder2': 'Live data for placeholder 2',
        '<<Figure4>>':'image',
        '<<Figure5>>':'image',
        '<<Figure6>>':'image',
        '<<Figure11>>':'image'

    }

    i=0
    for paragraph in doc.paragraphs:
        i+=1
        # print('PAGECOUNT',i)
        # print(paragraph)
        for placeholder, data in live_data.items():
            if placeholder in paragraph.text:
                run=paragraph.add_run()
                if placeholder=='<<Figure4>>':
                    plot_image_4=implement_unified_sport_activity(dashboard_filters,image=True)
                    run.add_picture(plot_image_4,width=width_,height=height_)
                    paragraph.text = paragraph.text.replace(placeholder, '')  #if you remove this ikage is getting removed
                # elif placeholder=='<<Figure5>>':
                #     plot_image_5=implement_youth_leadership_activity(dashboard_filters,image=True)
                #     run.add_picture(plot_image_5,width=width_,height=height_)
                #     # paragraph.text = paragraph.text.replace(placeholder, '')  #if you remove this ikage is getting removed
                # elif placeholder=='<<Figure6>>':
                #     plot_image_6=implement_school_engagement_activity(dashboard_filters,image=True)
                #     run.add_picture(plot_image_6,width=width_,height=height_)     
                #     # paragraph.text = paragraph.text.replace(placeholder, '')  #if you remove this ikage is getting removed     
                # elif placeholder=='<<Figure11>>':
                #     plot_image_11=sona_resources_useful(dashboard_filters,image=True)
                #     run.add_picture(plot_image_11,width=width_+2,height=height_)     
                #     # paragraph.text = paragraph.text.replace(placeholder, '')  #if you remove this ikage is getting removed         
                # else:
                #     paragraph.text = paragraph.text.replace(placeholder, data)



    return doc

def write_bytes_image(plot_image):
    with open('sample_f.png', 'wb') as f:
        f.write(plot_image.getvalue())


def write_word_doc(doc, filename):
    doc.save(filename)

# Fill placeholders in the template document
updated_doc = fill_placeholders()

# Write the updated Word document to a file

write_word_doc(updated_doc, f'../static/{state_full_name(STATE)}_document.docx')
